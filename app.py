from flask import Flask, render_template, request, redirect, url_for, session
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
import os
import pandas as pd
from datetime import date, datetime as dt
from datetime import datetime
from itsdangerous import URLSafeTimedSerializer, BadSignature
from sqlalchemy import or_
from sqlalchemy import and_ 
# from flask import Flask, render_template, request, redirect_url
import psycopg2
from datetime import date, datetime
import ast
import datetime
import os
import json

from flask_cors import CORS
import re
# import spacy
from flask_mail import Mail, Message
from flask import render_template, redirect, url_for, flash
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer
# from spacy.matcher import Matcher
from flask import Flask, request, render_template
from sqlalchemy import ForeignKey
from sqlalchemy.orm import relationship
from itsdangerous import URLSafeTimedSerializer
from flask import request, render_template, flash, redirect, url_for
import secrets
import secrets
from urllib.parse import quote_plus
from flask_migrate import Migrate
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from datetime import datetime
import pytz 

app = Flask(__name__)
cors = CORS(app)
# app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_SERVER'] = 'smtp.office365.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'ganesh.s@makonissoft.com'
app.config['MAIL_PASSWORD'] = 'Fol98135'
# app.config['MAIL_PASSWORD'] = 'tozvnmxbcejynxpe'
mail = Mail(app)

app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get("DATABASE_URL")

app.config['SECRET_KEY'] = secrets.token_hex(16)
engine = create_engine(app.config['SQLALCHEMY_DATABASE_URI'])
Session = sessionmaker(bind=engine)
app.config['SECRET_KEY'] = secrets.token_hex(16)
# Specify the folder where uploaded resumes will be stored
UPLOAD_FOLDER = 'C:/Users/Makonis/PycharmProjects/login/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
cors = CORS(app)
# Specify the allowed resume file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
db = SQLAlchemy(app)
migrate = Migrate(app, db)
from datetime import timedelta
#hello

# Specify the folder where uploaded resumes will be stored
# UPLOAD_FOLDER = 'static/'
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# cors = CORS(app)
# Specify the allowed resume file extensions
# ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}

# db = SQLAlchemy(app)
# migrate = Migrate(app, db)

def generate_verification_token(user_id):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    return serializer.dumps(user_id)

class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True,nullable=False)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(50), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    user_type = db.Column(db.String(20), nullable=False)
    client = db.Column(db.String(100))
    candidate = relationship("Candidate", back_populates="user", uselist=False)
    is_active = db.Column(db.Boolean, default=True)
    is_verified = db.Column(db.Boolean, default=False)
    created_by = db.Column(db.String(50))
    otp = db.Column(db.String(6), default=False)
    registration_completed = db.Column(db.String(50))
    filename = db.Column(db.String(100))
    # image_file = db.Column(db.String(1000))
    image_file=db.Column(db.LargeBinary)
    image_deleted=db.Column(db.Boolean, default=False)
    def serialize(self):
        return {
            'id': self.id,
            'username': self.username,
            'name': self.name,
            'email': self.email,
            'user_type': self.user_type,
            'client': self.client,
            'is_active': self.is_active,
            'is_verified': self.is_verified,
            'created_by': self.created_by,
            'otp': self.otp,
            'registration_completed': self.registration_completed
        }
        
class Candidate(db.Model):
    __tablename__ = 'candidates'
    id = db.Column(db.Integer, primary_key=True)
    job_id = db.Column(db.Integer)
    name = db.Column(db.String(100), nullable=False)
    mobile = db.Column(db.String(20), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    client = db.Column(db.String(100), nullable=False)
    current_company = db.Column(db.String(100))
    position = db.Column(db.String(100))
    profile = db.Column(db.String(200))
    current_job_location = db.Column(db.String(100))
    preferred_job_location = db.Column(db.String(100))
    # resume = db.Column(db.String(1000))
    resume = db.Column(db.LargeBinary)
    skills = db.Column(db.String(500))
    qualifications = db.Column(db.String(200))
    experience = db.Column(db.String(200))
    relevant_experience = db.Column(db.String(200))
    current_ctc = db.Column(db.String(200))
    expected_ctc = db.Column(db.String(200))
    notice_period = db.Column(db.String(20))
    last_working_date = db.Column(db.Date)
    buyout = db.Column(db.Boolean, default=False)
    holding_offer = db.Column(db.String(20))
    total = db.Column(db.Integer)
    package_in_lpa = db.Column(db.Float)
    recruiter = db.Column(db.String(100))
    management = db.Column(db.String(100))
    status = db.Column(db.String(100))
    reason_for_job_change = db.Column(db.String(200))
    remarks = db.Column(db.String(200))
    screening_done = db.Column(db.Boolean, default=False)
    rejected_at_screening = db.Column(db.Boolean, default=False)
    l1_cleared = db.Column(db.Boolean, default=False)
    rejected_at_l1 = db.Column(db.Boolean, default=False)
    dropped_after_clearing_l1 = db.Column(db.Boolean, default=False)
    l2_cleared = db.Column(db.Boolean, default=False)
    rejected_at_l2 = db.Column(db.Boolean, default=False)
    dropped_after_clearing_l2 = db.Column(db.Boolean, default=False)
    onboarded = db.Column(db.Boolean, default=False)
    dropped_after_onboarding = db.Column(db.Boolean, default=False)
    date_created = db.Column(db.Date, default=date.today)
    time_created = db.Column(db.Time, default=datetime.now().time)
    comments = db.Column(db.String(1000))
    linkedin_url = db.Column(db.String(1000))
    user_id = db.Column(db.Integer, ForeignKey('users.id'))
    period_of_notice = db.Column(db.String(1000))
    user = relationship("User", back_populates="candidate")
    reference = db.Column(db.String(200))
    reference_name = db.Column(db.String(200))
    reference_position = db.Column(db.String(200))
    reference_information = db.Column(db.String(200))
    data_updated_date = db.Column(db.Date)
    data_updated_time = db.Column(db.Time)

    def serialize(self):
        return {
            'id': self.id,
            'job_id': self.job_id,
            'name': self.name,
            'mobile': self.mobile,
            'email': self.email,
            'client': self.client,
            'current_company': self.current_company,
            'position': self.position,
            'profile': self.profile,
            'current_job_location': self.current_job_location,
            'preferred_job_location': self.preferred_job_location,
            'resume': self.resume,
            'skills': self.skills,
            'qualifications': self.qualifications,
            'experience': self.experience,
            'relevant_experience': self.relevant_experience,
            'current_ctc': self.current_ctc,
            'expected_ctc': self.expected_ctc,
            'notice_period': self.notice_period,
            'last_working_date': self.last_working_date.strftime('%Y-%m-%d') if self.last_working_date else None,
            'buyout': self.buyout,
            'holding_offer': self.holding_offer,
            'total': self.total,
            'package_in_lpa': self.package_in_lpa,
            'recruiter': self.recruiter,
            'management': self.management,
            'status': self.status,
            'reason_for_job_change': self.reason_for_job_change,
            'remarks': self.remarks,
            'screening_done': self.screening_done,
            'rejected_at_screening': self.rejected_at_screening,
            'l1_cleared': self.l1_cleared,
            'rejected_at_l1': self.rejected_at_l1,
            'dropped_after_clearing_l1': self.dropped_after_clearing_l1,
            'l2_cleared': self.l2_cleared,
            'rejected_at_l2': self.rejected_at_l2,
            'dropped_after_clearing_l2': self.dropped_after_clearing_l2,
            'onboarded': self.onboarded,
            'dropped_after_onboarding': self.dropped_after_onboarding,
            'date_created': self.date_created.strftime('%Y-%m-%d'),
            'time_created': self.time_created.strftime('%H:%M:%S'),
            'comments': self.comments,
            'linkedin_url': self.linkedin_url,
            'user_id': self.user_id,
            'period_of_notice': self.period_of_notice,
            'reference': self.reference,
            'reference_name': self.reference_name,
            'reference_position': self.reference_position,
            'reference_information': self.reference_information,
            
        }
        
class Career_user(db.Model):
    __tablename__ = 'career_users'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True, nullable=False)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(50), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    user_type = db.Column(db.String(50), default="career_visitor")


class Career_notification(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    recruiter_name = db.Column(db.String(100), nullable=False)
    notification_status = db.Column(db.Boolean, default=False)

    def __init__(self, recruiter_name, notification_status=False):
        self.recruiter_name = recruiter_name
        self.notification_status = notification_status


class JobPost(db.Model):
    __tablename__ = 'job_posts'

    id = db.Column(db.Integer, primary_key=True)
    client = db.Column(db.String(100))
    experience_min = db.Column(db.Integer)
    experience_max = db.Column(db.Integer)
    budget_min = db.Column(db.String(300))
    budget_max = db.Column(db.String(300))
    location = db.Column(db.String(100))
    shift_timings = db.Column(db.String(100))
    notice_period = db.Column(db.String(100))
    role = db.Column(db.String(100))
    detailed_jd = db.Column(db.Text)
    jd_pdf =  db.Column(db.String(1000))
    mode = db.Column(db.String(100))
    recruiter = db.Column(db.String(1000))
    management = db.Column(db.String(100))
    date_created = db.Column(db.Date)
    time_created = db.Column(db.Time)
    job_status = db.Column(db.String(20))
    job_type = db.Column(db.String(100))
    skills = db.Column(db.String(500))
    notification = db.Column(db.String(20))
    data_updated_date = db.Column(db.Date)
    data_updated_time = db.Column(db.Time)

    def __init__(self, client, experience_min, experience_max, budget_min, budget_max, location, shift_timings,
                 notice_period, role, detailed_jd,jd_pdf, mode, recruiter, management,job_status,job_type,skills):
        self.client = client
        self.experience_min = experience_min
        self.experience_max = experience_max
        self.budget_min = budget_min
        self.budget_max = budget_max
        self.location = location
        self.shift_timings = shift_timings
        self.notice_period = notice_period
        self.role = role
        self.detailed_jd = detailed_jd
        self.jd_pdf = jd_pdf
        self.mode = mode
        self.recruiter = recruiter
        self.management = management
        self.job_status = job_status
        self.job_type = job_type
        self.skills = skills

class Deletedcandidate(db.Model):
    _tablename_ = 'deletedcandidate'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), nullable=False)
    candidate_name = db.Column(db.String(100), nullable=False)
    candidate_email = db.Column(db.String(100), nullable=False)
    client = db.Column(db.String(100), nullable=False)
    profile = db.Column(db.String(100), nullable=False)
    status = db.Column(db.String(100), nullable=False)

# class Notification(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     recruiter_name = db.Column(db.String(100), nullable=False)
#     notification_status = db.Column(db.Boolean, default=False)

#     def __init__(self, recruiter_name, notification_status=False):
#         self.recruiter_name = recruiter_name
#         self.notification_status = notification_status
class Notification(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    job_post_id = db.Column(db.Integer, db.ForeignKey('job_posts.id'))
    recruiter_name = db.Column(db.String(100), nullable=False)
    notification_status = db.Column(db.Boolean, default=False)
    num_notification = db.Column(db.Integer, default=0)  # New column added

    def __init__(self, job_post_id, recruiter_name, notification_status=False):
        self.job_post_id = job_post_id
        self.recruiter_name = recruiter_name
        self.notification_status = notification_status
        self.num_notification = 0  # Default value for num_notification

@app.route('/check_candidate', methods=['POST'])
def check_candidate():
    clients = []
    profiles = []
    dates=[]
    job_ids=[]
    status=[]
    field = request.json['field']
    value = request.json['value']

    # Query the database to check for an existing candidate with the provided mobile or email
    existing_candidate = Candidate.query.filter(or_(Candidate.mobile == value, Candidate.email == value)).all()
    for i in existing_candidate:
        clients.append(" " + i.client + " ")
        profiles.append(" " + i.profile + " " )
        dates.append(i.date_created.strftime('%Y-%m-%d'))
        job_ids.append(i.job_id)
        status.append(i.status)

    # candidate = Candidate.query.filter_by(mobile=existing_candidate.mobile).first()
    if existing_candidate:
        response = {
            'message' : f"Candidate with this {field} already exists.",
            'client' : clients,
            'profile' : profiles,
            'dates':dates,
            'jobId':job_ids,
            'status':status
        }

    else:
        response = {
            'message': f"{field.capitalize()} is available.",
            'client': None,
            'profile': None,
            'dates':None,
            'jobId':None,
            'status':None
        }
    return json.dumps(response)

@app.route('/recruiter')
def recruiter_index():
    return render_template('recruiter_index.html')

@app.route('/')
def index():
    session_timeout_msg = request.args.get("session_timeout_msg")
    reset_message = request.args.get("reset_message")
    signup_message = request.args.get('signup_message')
    password_message = request.args.get('password_message')
    return render_template('index.html',reset_message=reset_message,session_timeout_msg=session_timeout_msg,signup_message=signup_message,password_message=password_message)

@app.route('/forgot_password')
def forgot_password():
    return render_template('forgot_password.html')

def generate_6otp():
    digits = "0123456789"
    otp = "".join(random.choice(digits) for _ in range(6))
    return otp


@app.route('/generate_otp', methods=['POST'])
def generate_otp():
    if request.method == 'POST':
        username = request.json.get('username')
        email = request.json.get('email')
        user = User.query.filter_by(username=username, email=email).first()
        if user:
            otp = generate_6otp()
            user.otp = otp
            db.session.commit()
            msg = Message('Account Verification', sender='ganesh.s@makonissoft.com', recipients=[email])
            msg.body = f'Hi {user.name},\n\n OTP for resetting your password {otp}.'
            mail.send(msg)
            return jsonify({'status': 'success', 'message': 'OTP has been sent to your email.'})
        else:
            return jsonify({'status': 'error', 'message': 'User does not exist.'})
    else:
        return jsonify({'status': 'error', 'message': 'Invalid request method.'})
    

   
import hashlib

@app.route('/reset_password', methods=['POST'])
def reset_password():
    if request.method == 'POST':
        otp = request.json['otp']
        new_password = request.json.get('new_password')
        confirm_password = request.json.get('confirm_password')
        new_password_hashed = hashlib.sha256(new_password.encode()).hexdigest()

        user = User.query.filter_by(otp=otp).first()

        if user and user.otp == otp and new_password == confirm_password:
            # Check if the new password is different from the old password
            if new_password_hashed != user.password:  # comparing hashes
                user.password = new_password_hashed
                db.session.commit()
                # Send the updated password to the user's email
                msg = Message('Password Changed', sender='ganesh.s@makonissoft.com', recipients=[user.email])
                msg.body = f'Hello {user.name},\n\nYour password has been successfully changed. Here are your updated credentials:\n\nUsername: {user.username}\nPassword: {new_password}'
                mail.send(msg)

                return jsonify({ 'message': 'Password changed successfully.'})
            else:
                return jsonify({ 'message': 'New password is the same as the old password'})
        else:
            return jsonify({ 'message': 'Invalid OTP or password confirmation. Please try again.'})

    return jsonify({ 'message': 'Invalid request method.'})

    

@app.route('/verify/<token>')
def verify(token):
    user_id = verify_token(token)
    if user_id:
        user = User.query.get(user_id)
        user.is_verified = True
        db.session.commit()
        if user.user_type == 'management':
            return jsonify({'status': 'success', 'message': 'Account verified successfully!', 'redirect': url_for('management_login', verification_msg_manager='Your Account has been Successfully Verified. Please Login.')})
        elif user.user_type == 'recruiter':
            return jsonify({'status': 'success', 'message': 'Account verified successfully!', 'redirect': url_for('recruiter_index')})
    else:
        return jsonify({'status': 'error', 'message': 'Your verification link has expired. Please contact management to activate your account.'})
    return jsonify({'status': 'error', 'message': 'An error occurred while verifying your account.'})


import hashlib
import random
import string

# Function to generate a random password
def generate_random_password(length=8):
    digits = string.digits
    password = ''.join(random.choice(digits) for _ in range(length - 3))
    return "Mak" + password

@app.route('/signup', methods=['POST'])
def signup():
    data = request.json
    user_id = data.get('user_id')  # Using get method to avoid KeyError
    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({'status': 'error', 'message': 'Invalid user ID or user does not exist.'})

    user_type = user.user_type
    user_name = user.username

    if user_type == 'management':
        username = data.get('username')
        name = data.get('name')
        email = data.get('email')
        user_type = data.get('user_type')

        # Check if required fields are provided
        if not all([username, name, email, user_type]):
            return jsonify({'status': 'error', 'message': 'All fields are required'})

        # Generate a random password
        password = generate_random_password()
        hashed_password = hashlib.sha256(password.encode()).hexdigest()

        created_by = user_name

        existing_user = User.query.filter(or_(User.username == username, User.email == email, User.name == name)).first()

        if existing_user:
            return jsonify({'status': 'error', 'message': 'Account with the same Username, Email, or Name already exists.'})

        new_user = User(username=username, password=hashed_password, name=name, email=email, user_type=user_type, created_by=created_by)
        
        db.session.add(new_user)
        db.session.commit()

        # Generate a verification token
        verification_token = generate_verification_token(new_user.id)

        # Create the verification link
        verification_link = url_for('verify', token=verification_token, _external=True)

        # Send the verification email
        msg = Message('Account Verification', sender='ganesh.s@makonissoft.com', recipients=[new_user.email])
        msg.body = f'Hello {new_user.name},\n\n We are pleased to inform you that your account has been successfully created for the ATS Makonis Talent Track Pro. Here are your login credentials:\n\nUsername: {new_user.username}\nPassword: {password}\n\n Please note that the verification link will expire after 24 hours. \n\n After successfully verifying your account, you can access the application using the following link : \n\n Application Link (Post Verification): https://ats-makonis.netlify.app/ \n\n To verify your account, please click on the following link: {verification_link} \n\n If you have any questions or need assistance, please feel free to reach out. \n\n Best regards, '
        mail.send(msg)

        # return jsonify({'message': 'A verification email has been sent to your email address. Please check your inbox.'})
        return jsonify({
            'message': 'A verification email has been sent to your email address. Please check your inbox.',
            'success_message': 'Account created successfully'
            })
    else:
        return jsonify({'message': 'You do not have permission to create recruiter accounts.'})


import hashlib

@app.route('/signup-onetime', methods=['POST'])
def signup_onetime():
    if request.method == 'POST':
        username = request.json.get('username')
        password = request.json.get('password')
        name = request.json.get('name')
        email = request.json.get('email')
        user_type = 'management'
        registration_completed = 'one_time'

        # Hash the password using SHA-256
        hashed_password = hashlib.sha256(password.encode()).hexdigest()

        user_onetime = User.query.filter_by(registration_completed='one_time').first()
        if user_onetime:
            return jsonify({'message': 'The one-time registration for this application has already been completed.'}),400

        new_user = User(username=username, password=hashed_password, name=name,
                        email=email, user_type=user_type, registration_completed=registration_completed)

        db.session.add(new_user)
        db.session.commit()

        # Generate a verification token
        verification_token = generate_verification_token(new_user.id)

        # Create the verification link
        verification_link = url_for('verify', token=verification_token, _external=True)

        # Construct the message body with username and plaintext password
        message_body = f'Hello {new_user.name},\n\nWe are pleased to inform you that your account has been successfully created for the ATS Makonis Talent Track Pro.\n\nYour login credentials:\n\nUsername: {new_user.username}\nPassword: {password}\n\nTo complete the account setup, kindly click on the verification link below:\n{verification_link}\n\nPlease note that the verification link will expire after 24 hours.\n\nAfter successfully verifying your account, you can access the application using the following link:\n\nApplication Link (Post Verification): https://ats-makonis.netlify.app/\n\nIf you have any questions or need assistance, please feel free to reach out.\n\nBest regards,'

        # Send the verification email
        msg = Message('Account Verification', sender='saiganeshkanuparthi@gmail.com', recipients=[new_user.email])
        msg.body = message_body
        mail.send(msg)

        # return jsonify({'message': 'A verification email has been sent to your email address. Please check your inbox.'})
        return jsonify({
            'message': 'A verification email has been sent to your email address. Please check your inbox.',
            'success_message': 'Account created successfully'
            }),200

    return jsonify({'message': 'Invalid request method.'}),400

@app.route('/login/recruiter', methods=['POST'])
def recruiter_login():
    verification_msg = request.args.get('verification_msg')
    reset_message = request.args.get('reset_message')
    session_timeout_msg = request.args.get("session_timeout_msg")
    password_message = request.args.get('password_message')

    if request.method == 'POST':
        username = request.json.get('username')
        password = request.json.get('password')

        # Hash the entered password
        hashed_password = hashlib.sha256(password.encode()).hexdigest()

        # Check if the user exists and the password is correct
        user = User.query.filter_by(username=username, password=hashed_password, user_type='recruiter').first()

        if user:
            if user.is_active:  # Check if the user is active
                if user.is_verified:
                    # Set the user session variables
                    session['user_id'] = user.id
                    session['user_type'] = user.user_type
                    session['username'] = user.username
                    session['user_name'] = user.name
                    session['JWT Token'] = secrets.token_hex(16)
                    return jsonify({'status': 'success', 'redirect': url_for('dashboard'),'user_id': user.id})
                else:
                    message = 'Your account is not verified yet. Please check your email for the verification link.'
            else:
                message = 'Your account is not active. Please contact the administrator.'
        else:
            message = 'Invalid username or password'

        return jsonify({'status': 'error', 'message': message})

    # For GET requests, return necessary data
    return jsonify({
        'status': 'success',
        'verification_msg': verification_msg,
        'reset_message': reset_message,
        'session_timeout_msg': session_timeout_msg,
        'password_message': password_message
    })


import hashlib

import hashlib

@app.route('/login/management', methods=['POST'])
def management_login():
    username = request.json.get('username')
    password = request.json.get('password')
    verification_msg_manager = request.args.get('verification_msg_manager')
    
    # Check if the user exists
    user = User.query.filter_by(username=username, user_type='management').first()
    
    if user:
        # Hash the provided password using the same hash function and parameters used to hash the passwords in the database
        hashed_password = hashlib.sha256(password.encode()).hexdigest()
        
        # Compare the hashed password with the hashed password stored in the database
        if hashed_password == user.password:
            if user.is_active:  # Check if the user is active
                if user.is_verified:
                    # Set the user session variables
                    session['user_id'] = user.id
                    session['user_type'] = user.user_type
                    session['username'] = user.username
                    session['user_name'] = user.name
                    session['JWT Token'] = secrets.token_hex(16)
                    return jsonify({'status': 'success', 'redirect': url_for('dashboard'),'user_id':user.id})
                else:
                    message = 'Your account is not verified yet. Please check your email for the verification link.'
            else:
                message = 'Your account is not active. Please contact the administrator.'
        else:
            message = 'Invalid username or password'
    else:
        message = 'Invalid username or password'

    return jsonify({'status': 'error', 'message': message, 'verification_msg_manager': verification_msg_manager})

# @app.route('/get_recruiters', methods=['GET'])   
# def get_recruiters_list():
#     recruiters = User.query.filter_by(user_type='recruiter').all()
    
#     # Assuming you want to return a list of dictionaries containing user details
#     recruiters_list = []
#     for recruiter in recruiters:
#         recruiter_dict = {
#             'id': recruiter.id,
#             'username': recruiter.username,
#             'user_type': recruiter.user_type
#             # Add more fields if needed
#         }
#         recruiters_list.append(recruiter_dict)
    
#     return jsonify(recruiters_list)


@app.route('/get_recruiters', methods=['GET'])   
def get_recruiters_list():
    recruiters = User.query.filter_by(user_type='recruiter').all()
    
    # Extracting only usernames
    usernames = [recruiter.username for recruiter in recruiters]
    
    return jsonify(usernames)

@app.route('/get_recruiters_candidate', methods=['POST'])
def recruiter_candidate_list():
    data = request.json
    
    if not data or 'user_name' not in data:
        return jsonify({'error': 'Invalid input'}), 400
    
    username = data['user_name']
    
    # Find the recruiter with the given username
    recruiter = User.query.filter_by(username=username, user_type='recruiter').first()
    print("recruiter:",recruiter)
    
    if recruiter:
        # Find all candidates linked with the recruiter's username
        candidates = Candidate.query.filter_by(recruiter=username).all()
        
        # Prepare response data
        candidates_list = [
            {
                'id': candidate.id,
                'username': candidate.name,
                'status': candidate.status,
                'profile': candidate.profile,
                'recruiter': candidate.recruiter
            } 
            for candidate in candidates
        ]
        return jsonify(candidates_list)
    else:
        return jsonify({'error': 'Recruiter not found'}), 404

# @app.route('/get_recruiters_candidate', methods=['POST']) 
# def recruiter_candidate_list():
#     data = request.json
#     username = data['user_name']
    
#     # Find the recruiter with the given username
#     recruiter = User.query.filter_by(username=username, user_type='recruiter').first()
    
#     if recruiter:
#         # Find all candidates linked with the recruiter's username
#         candidates = Candidate.query.filter_by(recruiter=recruiter.username).all()
        
#         # Prepare response data
#         candidates_list = [{'id': candidate.id, 'username': candidate.name, 'status':candidate.status, 'profile':candidate.profile, 'recruiter':candidate.recruiter} for candidate in candidates]
#         return jsonify(candidates_list)
#     else:
#         return jsonify({'error': 'Recruiter not found'})


@app.route('/assign_candidate_new_recuriter', methods=['POST']) 
def assign_candidate_to_a_new_recruiter():
    data = request.json

    try:
        candidates_data = []
        for candidate_data in data['candidates']:
            candidate_id = candidate_data.get('candidate_id')
            new_recruiter_username = candidate_data.get('new_recruiter')
            current_recruiter_username = candidate_data.get('current_recruiter')

            if not candidate_id or not new_recruiter_username or not current_recruiter_username:
                return jsonify({"error": "Candidate ID, new recruiter username, or current recruiter username not provided"}), 400

            # Get the candidate, current recruiter, and the new recruiter from the database using their usernames
            candidate = Candidate.query.filter_by(id=candidate_id, recruiter=current_recruiter_username).first()
            # if candidate.profile_transfered != None:
            #     candidate.profile_transfered = "YES" 
            # else:
            #     candidate.profile_transfered = None
            current_recruiter = User.query.filter_by(username=current_recruiter_username, user_type='recruiter').first()
            new_recruiter = User.query.filter_by(username=new_recruiter_username, user_type='recruiter').first()

            if candidate is None:
                return jsonify({"error": "Candidate not found or not assigned to current recruiter"}), 404

            if current_recruiter is None:
                return jsonify({"error": "Current recruiter not found or not a recruiter"}), 404

            if new_recruiter is None:
                return jsonify({"error": "New recruiter not found or not a recruiter"}), 404

            # Update the candidate record to point to the new recruiter
            candidate.recruiter = new_recruiter_username
            db.session.commit()

            candidates_data.append({'id': candidate.id, 'name': candidate.name})

        return jsonify({
            "message": "Candidates assigned successfully.",
            "candidates": candidates_data
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": "Error assigning candidates: " + str(e)}), 500


from flask import jsonify

@app.route('/candidate_details/<int:candidate_id>/<user_type>/<int:page_no>', methods=['GET'])
def candidate_details(candidate_id, user_type, page_no):
    user_name = session.get('user_name')
    count_notification_no = Notification.query.filter(Notification.notification_status == 'false', Notification.recruiter_name == user_name).count()
    candidate = Candidate.query.get(candidate_id)
    if candidate:
        # Return JSON response with candidate details
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "mobile": candidate.mobile,
            "email": candidate.email,
            "client": candidate.client,
            "current_company": candidate.current_company,
            "position": candidate.position,
            "profile": candidate.profile,
            "current_job_location": candidate.current_job_location,
            "preferred_job_location": candidate.preferred_job_location,
            "resume": candidate.resume,
            "skills": candidate.skills,
            "qualifications": candidate.qualifications,
            "experience": candidate.experience,
            "relevant_experience": candidate.relevant_experience,
            "current_ctc": candidate.current_ctc,
            "expected_ctc": candidate.expected_ctc,
            "linkedin_url": candidate.linkedin_url,
            "notice_period": candidate.notice_period,
            "holding_offer": candidate.holding_offer,
            "user_type": user_type,
            "user_name": user_name,
            "count_notification_no": count_notification_no,
            "page_no": page_no
        })
    else:
        # Return JSON response with error message
        return jsonify({"error_message": "Candidate not found"}), 404


from flask import Flask, jsonify, request, Response

@app.route('/dashboard', methods=['POST'])
def dashboard():
    data = request.json
    print(data)  # Just to verify if data is received properly

    edit_candidate_message = data.get('edit_candidate_message')
    page_no = data.get('page_no')
    candidate_message = data.get('candidate_message')
    signup_message = data.get('signup_message')
    job_message = data.get('job_message')
    update_candidate_message = data.get('update_candidate_message')
    delete_message = data.get("delete_message")

    # data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()
    user_type = user.user_type
    user_name = user.username

    response_data = {}

    if user_id and user_type:
        if user_type == 'recruiter':
            recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
            if recruiter:
                candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name, Candidate.reference.is_(None))).all()  # Filter candidates by recruiter's name
                candidates = sorted(candidates, key=lambda candidate: candidate.id)
                jobs = JobPost.query.filter_by(recruiter=user_name).all()  # Filter jobs by recruiter's name
                count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                                  Notification.recruiter_name == user_name).count()
                career_count_notification_no = Career_notification.query.filter(Career_notification.notification_status == 'false',
                                                                  Career_notification.recruiter_name == user_name).count()
                response_data = {
                    'user': {
                        'id': recruiter.id,
                        'name': recruiter.name,
                        'user_type': recruiter.user_type,
                        'email': recruiter.email
                        # Add more attributes as needed
                    },
                    'user_type': user_type,
                    'user_name': user_name,
                    'candidates': [{
                        'id': candidate.id,
                        'job_id': candidate.job_id,
                        'name': candidate.name,
                        'mobile': candidate.mobile,
                        'email': candidate.email,
                        'client': candidate.client,
                        'current_company': candidate.current_company,
                        'position': candidate.position,
                        'profile': candidate.profile,
                        'current_job_location': candidate.current_job_location,
                        'preferred_job_location': candidate.preferred_job_location,
                        'qualifications': candidate.qualifications,
                        'experience': candidate.experience,
                        'relevant_experience': candidate.relevant_experience,
                        'current_ctc': candidate.current_ctc,
                        'expected_ctc': candidate.expected_ctc,
                        'notice_period': candidate.notice_period,
                        'linkedin_url': candidate.linkedin_url,
                        'holding_offer': candidate.holding_offer,
                        'recruiter': candidate.recruiter,
                        'management': candidate.management,
                        'status': candidate.status,
                        'remarks': candidate.remarks,
                        'skills': candidate.skills,
                        'resume': candidate.resume,
                        'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
                        'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
                        'buyout': candidate.buyout,
                        'date_created':candidate.date_created,
                        'time_created':candidate.time_created

            
                        # Add more attributes as needed
                    } for candidate in candidates],
                    'jobs': [{
                        'id': job.id,
                        'client': job.client,
                        'experience_min': job.experience_min,
                        'experience_max': job.experience_max,
                        'budget_min': job.budget_min,
                        'budget_max': job.budget_max,
                        'location': job.location,
                        'shift_timings': job.shift_timings,
                        'notice_period': job.notice_period,
                        'role': job.role,
                        'detailed_jd': job.detailed_jd,
                        'jd_pdf': job.jd_pdf,
                        'mode': job.mode,
                        'recruiter': job.recruiter,
                        'management': job.management,
                        'date_created': job.date_created,
                        'time_created': job.time_created,
                        'job_status': job.job_status,
                        'job_type': job.job_type,
                        'skills': job.skills,
                        'notification': job.notification
                        # Add more attributes as needed
                    } for job in jobs],
                    'edit_candidate_message': edit_candidate_message,
                    'page_no': page_no,
                    'count_notification_no': count_notification_no,
                    'career_count_notification_no': career_count_notification_no
                }
        elif user_type == 'management':
            users = User.query.all()
            candidates = Candidate.query.filter(Candidate.reference.is_(None)).all()
            candidates = sorted(candidates, key=lambda candidate: candidate.id)
            jobs = JobPost.query.all()
            response_data = {
                'users': [{
                    'id': user.id,
                    'name': user.name,
                    'user_type': user.user_type,
                    'email': user.email
                     
                    # Add more attributes as needed
                } for user in users],
                'user_type': user_type,
                'user_name': user_name,
                'candidates': [{
                        'id': candidate.id,
                        'job_id': candidate.job_id,
                        'name': candidate.name,
                        'mobile': candidate.mobile,
                        'email': candidate.email,
                        'client': candidate.client,
                        'current_company': candidate.current_company,
                        'position': candidate.position,
                        'profile': candidate.profile,
                        'current_job_location': candidate.current_job_location,
                        'preferred_job_location': candidate.preferred_job_location,
                        'qualifications': candidate.qualifications,
                        'experience': candidate.experience,
                        'relevant_experience': candidate.relevant_experience,
                        'current_ctc': candidate.current_ctc,
                        'expected_ctc': candidate.expected_ctc,
                        'notice_period': candidate.notice_period,
                        'linkedin_url': candidate.linkedin_url,
                        'holding_offer': candidate.holding_offer,
                        'recruiter': candidate.recruiter,
                        'management': candidate.management,
                        'status': candidate.status,
                        'remarks': candidate.remarks,
                        'skills': candidate.skills,
                        'resume': candidate.resume,
                        'period_of_notice': candidate.period_of_notice if candidate.notice_period == 'no' else None,
                        'last_working_date': candidate.last_working_date if candidate.notice_period in {'yes', 'completed'} else None,
                        'buyout': candidate.buyout,
                        'date_created':candidate.date_created,
                        'time_created':candidate.time_created

                    # Add more attributes as needed
                } for candidate in candidates],
                'jobs': [{
                    'id': job.id,
                    'client': job.client,
                    'experience_min': job.experience_min,
                    'experience_max': job.experience_max,
                    'budget_min': job.budget_min,
                    'budget_max': job.budget_max,
                    'location': job.location,
                    'shift_timings': job.shift_timings,
                    'notice_period': job.notice_period,
                    'role': job.role,
                    'detailed_jd': job.detailed_jd,
                    'jd_pdf': job.jd_pdf,
                    'mode': job.mode,
                    'recruiter': job.recruiter,
                    'management': job.management,
                    'date_created': job.date_created,
                    'time_created': job.time_created,
                    'job_status': job.job_status,
                    'job_type': job.job_type,
                    'skills': job.skills,
                    'notification': job.notification
                    # Add more attributes as needed
                } for job in jobs],
                'signup_message': signup_message,
                'job_message': job_message,
                'page_no': page_no,
                'edit_candidate_message': edit_candidate_message
            }
        else:
            user = User.query.filter_by(id=user_id).first()
            if user:
                candidates = Candidate.query.filter_by(recruiter=user.name).all()  # Filter candidates by user's name
                response_data = {
                    'user': {
                        'id': user.id,
                        'name': user.name,
                        'user_type': user.user_type,
                        'email': user.email
                        # Add more attributes as needed
                    },
                    'user_type': user_type,
                    'user_name': user_name,
                    'candidates': [{
                        'id': candidate.id,
                        'job_id':candidate.job_id,
                        'name': candidate.name,
                        'email': candidate.email,
                        'mobile': candidate.mobile,
                        'client':candidate.client,
                        'skills':candidate.skills,
                        "profile": candidate.profile, 
                        'recruiter':candidate.recruiter,
                        "management":candidate.management,
                        'resume': candidate.resume,
                        'current_company': candidate.current_company,
                        'position': candidate.position,
                        'current_job_location': candidate.current_job_location,
                        'preferred_job_location': candidate.preferred_job_location,
                        'qualifications':candidate.qualifications,
                        'experience': candidate.experience,
                        'relevant_experience':candidate.relevant_experience,
                        'current_ctc':candidate.current_ctc,
                        'experted_ctc': candidate.expected_ctc,
                        "total":candidate.total,
                        'package_in_lpa':candidate.package_in_lpa,
                        'holding_offer':candidate.holding_offer,
                        'status': candidate.status,
                        'reason_for_job_change':candidate.reason_for_job_change,
                        'remarks':candidate.remarks,
                        'screening_done': candidate.screening_done,
                        'rejected_at_screening': candidate.rejected_at_screening,
                        'l1_cleared':candidate.l1_cleared,
                        'rejected_at_l1':candidate.rejected_at_l1,
                        "dropped_after_clearing_l1": candidate.dropped_after_clearing_l1,
                        'l2_cleared':candidate.l1_cleared,
                        'rejected_at_l2':candidate.rejected_at_l1,
                        "dropped_after_clearing_l2": candidate.dropped_after_clearing_l1,
                        'onboarded': candidate.onboarded,
                        'dropped_after_onboarding': candidate.dropped_after_onboarding,
                        'linkedin_url': candidate.linkedin_url,
                        'period_of_notice': candidate.period_of_notice,
                        'reference': candidate.reference,
                        'reference_name': candidate.reference_name,
                        'reference_position': candidate.reference_position,
                        'reference_information': candidate.reference_information,
                        'comments':candidate.comments,
                        "time_created":str(candidate.time_created),
                        "date_created": str(candidate.date_created)
                        # Add more attributes as needed
                    } for candidate in candidates],
                }
    else:
        response_data = {"message": "User ID or User Type missing"}

    # Convert date objects to string representations before returning the response
    for job in response_data.get('jobs', []):
        job['date_created'] = job['date_created'].isoformat()

    return Response(json.dumps(response_data, default=str), content_type='application/json')


# Mocked function for demonstration
# Mocked function for demonstration
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'doc', 'docx'}


import binascii  

import binascii  

@app.route('/add_candidate', methods=['POST'])
def add_candidate():
    try:
        
        # Retrieve request data from JSON
        data = request.json
        user_id = data['user_id']
        user = User.query.filter_by(id=user_id).first()
        user_type = user.user_type
        user_name = user.username

        job_id = data.get('job_id')
        client = data.get('client')
        name = data.get('name')
        mobile = data.get('mobile')
        email = data.get('email')
        profile = data.get('profile')
        skills = data.get('skills')
        current_company = data.get('current_company')
        position = data.get('position')
        current_job_location = data.get('current_job_location')
        preferred_job_location = data.get('preferred_job_location')
        qualifications = data.get('qualifications')
        experience = data.get('experience')
        experience_months=data.get('experience')
        relevant_experience = data.get('relevant_experience')
        relevant_experience_months=data.get('relevant_experience_months')
        reason_for_job_change=data.get('reason_for_job_change')
        current_ctc = data.get('current_ctc')
        expected_ctc = data.get('expected_ctc')
        linkedin = data.get('linkedin')
        notice_period = data.get('notice_period')
        holding_offer = data.get('holding_offer')
        resume = data.get('resume')
        print("Resume : ",type(resume))


        

        # # Check if the user is logged in
        if request.method == 'POST':
              
            # Retrieve the recruiter and management names based on user type
            if user_type == 'recruiter':
                recruiter = User.query.get(user_id).name
                management = None
            elif user_type == 'management':
                recruiter = None
                management = User.query.get(user_id).name
            else:
                recruiter = None
                management = None

            # Check if the job_id is provided and job is active
            matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Active')).first()
            if not matching_job_post:
                return jsonify({"error_message": "Job on hold"})

            # Create new candidate object
            new_candidate = Candidate(
                user_id=user_id,
                job_id=job_id,
                name=name,
                mobile=mobile,
                email=email,
                client=client,
                current_company=current_company,
                position=position,
                profile=profile,
                current_job_location=current_job_location,
                preferred_job_location=preferred_job_location,
                qualifications=qualifications,
                experience=experience,
                relevant_experience=relevant_experience,
                current_ctc=current_ctc,
                expected_ctc=expected_ctc,
                notice_period=notice_period,
                linkedin_url=linkedin,
                holding_offer=holding_offer,
                recruiter=recruiter,
                management=management,
                status='None',
                remarks=data.get('remarks'),
                skills=skills,
                resume=resume,
                period_of_notice=data.get('months') if notice_period == 'no' else None,
                last_working_date=data.get('last_working_date') if notice_period in {'yes', 'completed'} else None,
                buyout='buyout' in data
            )

            new_candidate.date_created = date.today()
            new_candidate.time_created = datetime.now().time()

            db.session.add(new_candidate)
            db.session.commit()

            return jsonify({"message": "Candidate Added Successfully", "candidate_id": new_candidate.id})

        return jsonify({"error_message": "Method not found"})

    except Exception as e:
        return jsonify({"error_message": str(e)}),500
        
        
from flask import jsonify

@app.route('/get_job_role', methods=['GET'])
def get_job_role():
    job_id = request.args.get('job_id')

    job_post = JobPost.query.filter_by(id=job_id).first()
    if job_post:
        return jsonify({"role": job_post.role})
    else:
        return jsonify({"role": ""})

@app.route('/delete_candidate/<int:candidate_id>', methods=["POST"])
def delete_candidate(candidate_id):
    data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()
    user_type = user.user_type
    username=user.username
    
    if user_type == 'management':
        candidate = Candidate.query.filter_by(id=candidate_id).first()

        if candidate:
            if request.method == "POST":
                # Save deletion details before deleting the candidate
                deleted_candidate = Deletedcandidate(
                    username=username,
                    candidate_name=candidate.name,
                    candidate_email=candidate.email,
                    client=candidate.client,
                    profile=candidate.profile,
                    status=candidate.status
                )
                db.session.add(deleted_candidate)
                db.session.commit()

                # Delete the candidate
                Candidate.query.filter_by(id=candidate_id).delete()
                db.session.commit()

                return jsonify({"message": "Candidate details deleted successfully"})

            return jsonify({
                "candidate": {
                    "id": candidate.id,
                    "name": candidate.name,
                    "email": candidate.email,
                    "client": candidate.client,
                    "profile": candidate.profile,
                    "status": candidate.status
                },
                "user_name": username
            })

        else:
            return jsonify({"message": "Candidate not found"}), 404

    return jsonify({"message": "Unauthorized: Only management can delete candidates"}), 401


@app.route('/delete_candidate_recruiter/<int:candidate_id>', methods=["GET", "POST"])
def delete_candidate_recruiter(candidate_id):
    if 'user_id' in session and 'user_type' in session:
        user_type = session['user_type']
        user_name = session['user_name']

        if user_type == 'management':
            candidate = Candidate.query.filter_by(id=candidate_id,recruiter=user_name)

            if request.method == "POST":
                # Save deletion details before deleting the candidate
                deleted_candidate = Deletedcandidate(
                    username=user_name,
                    candidate_name=candidate.name,
                    candidate_email=candidate.email,
                    client=candidate.client,
                    profile=candidate.profile,
                    status=candidate.status
                )
                db.session.add(deleted_candidate)
                db.session.commit()

                # Delete the candidate
                Candidate.query.filter_by(id=candidate_id).delete()
                db.session.commit()

                return redirect(url_for('dashboard', delete_message="Candidate details deleted successfully"))

            return render_template('delete_candidate.html', candidate=candidate, user_name=user_name)

        return "Unauthorized: Only management can delete candidates", 401

    return "Unauthorized: You must log in to access this page", 401


def verify_token(token):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    try:
        user_id = serializer.loads(token, max_age=86400)  
        return user_id
    except BadSignature:
        return None  
    except Exception as e:
        return None

# Search String Changed
# @app.route('/update_candidate/<int:candidate_id>/<page_no>/<search_string>', methods=['GET', 'POST'])
from flask import session, jsonify

@app.route('/update_candidate/<int:candidate_id>', methods=['POST'])
def update_candidate(candidate_id):
    data = request.json

    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()
    user_type = user.user_type
    user_name = user.username
    count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                      Notification.recruiter_name == user_name).count()
    career_count_notification_no = Career_notification.query.filter(
        Career_notification.notification_status == 'false',
        Career_notification.recruiter_name == user_name).count()
    if request.method == 'POST':
        if user_type == 'recruiter':
            recruiter = User.query.get(user_id).name
            management = None
        elif user_type == 'management':
            recruiter = None
            management = User.query.get(user_id).name
        else:
            recruiter = None
            management = None

        if user_type == 'recruiter':
            user_email = User.query.get(user_id).email
            management_email = None
        elif user_type == 'management':
            user_email = None
            management_email = User.query.get(user_id).email
        else:
            user_email = None
            management_email = None

        candidate = Candidate.query.filter_by(id=candidate_id).first()
        print(candidate)
        
        previous_status = candidate.status

        candidate_status = request.json.get('candidate_status')
        candidate_comment = request.json.get('comments')

        candidate.status = candidate_status
        candidate.comments = candidate_comment

        db.session.commit()

        if candidate_status in [
                "SCREENING", "SCREEN REJECTED", "NO SHOW", "DROP", "CANDIDATE HOLD", "OFFERED - DECLINED", "DUPLICATE", "SCREENING SELECTED",
                "L1-SCHEDULE", "L1-FEEDBACK", "L1-SELECTED", "L1-REJECTED", "CANDIDATE RESCHEDULE", "PANEL RESCHEDULE", "L2-SCHEDULE", 
                "L2-FEEDBACK", "L2-SELECTED", "L2-REJECTED", "HR-ROUND", "MANAGERIAL ROUND", "NEGOTIATION", "SELECTED", "OFFER-REJECTED",
                "OFFER-DECLINED", "ON-BOARDED", "HOLD", "CANDIDATE NO-SHOW"
                ]:
            candidate_name = candidate.name
            candidate_position = candidate.position
            candidate_email = candidate.email

            if candidate_position:
                candidate_position = candidate_position.upper()
            else:
                candidate_position = ""

            if candidate.client:
                client = candidate.client.upper()
            else:
                client = ""

            if candidate_status in ["SCREENING", "SCREEN REJECTED"]:
                message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
            else:
                message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Previous Status : "{previous_status}"\n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
        else:
            message = ""
            candidate_name = ""
            candidate_position = ""
            candidate_email = ""

        return jsonify({
        "message": "Candidate Status Updated Successfully",
        "user_id": user_id,
        "user_type": user_type,
        "user_name": user_name,
        "count_notification_no": count_notification_no,
        "career_count_notification_no": career_count_notification_no,
        "recruiter": recruiter,
        "management": management,
        "recruiter_email": user_email,
        "management_email": management_email,
        "candidate_name": candidate_name,
        "candidate_position": candidate_position,
        "candidate_email": candidate_email,
        # "message": message
        "message_body": message 
    })

@app.route('/update_candidate_careers/<int:candidate_id>/<page_no>/<search_string>', methods=['GET', 'POST'])
@app.route('/update_candidate_careers/<int:candidate_id>/<page_no>', methods=['GET', 'POST'])
def update_candidate_careers(candidate_id, page_no):
    if 'user_id' in session and 'user_type' in session:
        user_id = session['user_id']
        user_type = session['user_type']
        user_name = session['user_name']
        count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                          Notification.recruiter_name == user_name).count()
        career_count_notification_no = Career_notification.query.filter(
            Career_notification.notification_status == 'false',
            Career_notification.recruiter_name == user_name).count()
        if request.method == 'POST':
            # Retrieve the logged-in user's ID and user type from the session
            user_id = session['user_id']
            user_type = session['user_type']

            # Retrieve the recruiter and management names based on user type
            if user_type == 'recruiter':
                recruiter = User.query.get(user_id).name
                management = None
            elif user_type == 'management':
                recruiter = None
                management = User.query.get(user_id).name
            else:
                recruiter = None
                management = None

            if user_type == 'recruiter':
                recruiter_email = User.query.get(user_id).email
                management = None
            elif user_type == 'management':
                recruiter = None
                recruiter_email = User.query.get(user_id).email
            else:
                recruiter_email = None
                management_email = None

            # Retrieve the form data for the candidate
            candidate = Candidate.query.get(candidate_id)
            previous_status = candidate.status
            # candidate.recruiter = recruiter
            # candidate.management = management

            # Get the selected candidate status from the form
            candidate_status = request.form.get('candidate_status')
            candidate_comment = request.form.get('comments')

            # Update the candidate status field
            candidate.status = candidate_status

            candidate.comments = candidate_comment

            db.session.commit()

            if candidate_status == "SCREENING" or candidate_status == "SCREEN REJECTED":
                candidate_name = candidate.name
                candidate_position = candidate.position

                # Retrieve the candidate's email
                candidate_email = candidate.email

                # Determine if the logged-in user is a recruiter or management
                user_type = session.get('user_type')

                if user_type == 'recruiter' or user_type == 'management':
                    # Retrieve the corresponding user's email
                    user_email = User.query.get(session.get('user_id')).email

                    message = Message(f'Job Application Status - {candidate_position}',
                                      sender='ganesh.s@makonissoft.com', recipients=[candidate_email])

                    if user_type == 'management':
                        management_email = user_email
                        message.cc = [management_email]
                    elif user_type == 'recruiter':
                        recruiter_email = user_email
                        message.cc = [recruiter_email]
                    message.body = f'''Dear {candidate.name}, 

Greetings! 

We hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate.position.upper()} position and participating in the recruitment process. 

We are writing to inform you about the latest update we received from our client {candidate.client.upper()} regarding your interview. 

        Current Status :  "{candidate.status}"

Thank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. 

If you have any questions or need further information, please feel free to reach out to us. 

Thanks, 
                            '''
                #mail.send(message)
                pass
            elif candidate_status == "NO SHOW" or candidate_status == "DROP" or candidate_status == "CANDIDATE HOLD" or candidate_status == "OFFERED - DECLINED" or candidate_status == "DUPLICATE":
                pass
            else:
                candidate_name = candidate.name
                candidate_position = candidate.position
                candidate_email = candidate.email

                user_type = session.get('user_type')

                if user_type == 'recruiter' or user_type == 'management':
                    user_email = User.query.get(session.get('user_id')).email

                    message = Message(f'Job Application Status - {candidate_position}',
                                      sender='ganesh.s@makonissoft.com', recipients=[candidate_email])

                    if user_type == 'management':
                        management_email = user_email
                        message.cc = [management_email]
                    elif user_type == 'recruiter':
                        recruiter_email = user_email
                        message.cc = [recruiter_email]
                    message.body = f'''Dear {candidate.name}, 

Greetings! 

We hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate.position.upper()} position and participating in the recruitment process. 

We are writing to inform you about the latest update we received from our client {candidate.client.upper()} regarding your interview. 

        Previous Status : "{previous_status}"

        Current Status :  "{candidate.status}"

Thank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. 

If you have any questions or need further information, please feel free to reach out to us. 

Thanks, 
                            '''
                #mail.send(message)
                pass

            return redirect(
                url_for('career_dashboard', update_candidate_message='Candidate Status Updated Sucessfully', page_no=page_no))

        candidate = Candidate.query.get(candidate_id)
        candidate_data = {
            'id': candidate.id,
            'name': candidate.name,
            'mobile': candidate.mobile,
            'email': candidate.email,
            'client': candidate.client,
            'current_company': candidate.current_company,
            'position': candidate.position,
            'profile': candidate.profile,
            'current_job_location': candidate.current_job_location,
            'preferred_job_location': candidate.preferred_job_location,
            'resume': candidate.resume,
            'qualifications': candidate.qualifications,
            'experience': candidate.experience,
            'relevant_experience': candidate.relevant_experience,
            'current_ctc': candidate.current_ctc,
            'expected_ctc': candidate.expected_ctc,
            'notice_period': candidate.notice_period,
            'last_working_date': candidate.last_working_date,
            'buyout': candidate.buyout,
            'holding_offer': candidate.holding_offer,
            'total': candidate.total,
            'package_in_lpa': candidate.package_in_lpa,
            'reason_for_job_change': candidate.reason_for_job_change,
            'remarks': candidate.remarks,
            'candidate_status': candidate.status,
        }

        return render_template('update_candidate.html', candidate_data=candidate_data, user_id=user_id,
                               user_type=user_type, user_name=user_name, candidate=candidate,
                               count_notification_no=count_notification_no,
                               career_count_notification_no=career_count_notification_no)

    return redirect(url_for('career_dashboard'))


from flask import jsonify

@app.route('/logout', methods=['POST'])
def logout():
    data = request.json
    
    if data:
        user_id = data.get('user_id')
        
        if user_id:
            user = User.query.filter_by(id=user_id).first()
            
            if user:
                user_type = user.user_type
                user_name = user.username
                
                return jsonify({"message": "Logged out successfully"}), 200
            
            return jsonify({"message": "User not found"}), 404
        else:
            return jsonify({"message": "'user_id' not provided in JSON data"}), 400
    
    return jsonify({"message": "No JSON data provided"}), 400



from datetime import datetime

# Search String Changed
# @app.route('/edit_candidate/<int:candidate_id>/<int:page_no>/<search_string>', methods=['GET', 'POST'])
@app.route('/edit_candidate/<int:candidate_id>', methods=['POST'])
def edit_candidate(candidate_id):
        data = request.json
        user_id = data['user_id']
        user = User.query.filter_by(id=user_id).first()
        user_name = user.username
        count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                          Notification.recruiter_name == user_name).count()
        career_count_notification_no = Career_notification.query.filter(
            Career_notification.notification_status == 'false',
            Career_notification.recruiter_name == user_name).count()

        if request.method == 'POST':
            # Retrieve the form data for the candidate from JSON payload
            data = request.json

            # Retrieve the candidate object
            candidate = Candidate.query.get(candidate_id)
            if candidate:
                # Update the candidate fields with the new data
                candidate.name = data.get('name')
                candidate.mobile = data.get('mobile')
                candidate.email = data.get('email')
                candidate.client = data.get('client')
                candidate.current_company = data.get('current_company')
                candidate.position = data.get('position')
                candidate.profile = data.get('profile')
                candidate.current_job_location = data.get('current_job_location')
                candidate.preferred_job_location = data.get('preferred_job_location')
                candidate.qualifications = data.get('qualifications')
                candidate.experience = data.get('experience')
                candidate.notice_period = data.get('notice_period')
                candidate.reason_for_job_change = data.get('reason_for_job_change')
                candidate.linkedin_url = data.get('linkedin')
                candidate.remarks = data.get('remarks')
                candidate.skills = data.get('skills')
                candidate.holding_offer = data.get('holding_offer')
                candidate.total = data.get('total')
                candidate.package_in_lpa = data.get('package_in_lpa')
                candidate.period_of_notice = data.get('period_of_notice')

                db.session.commit()
                return jsonify({"message": "Candidate Details Edited Successfully"})
            else:
                return jsonify({"error_message": "Candidate not found"}), 500


@app.route('/edit_candidate_careers/<int:candidate_id>/<int:page_no>/<search_string>', methods=['GET', 'POST'])
@app.route('/edit_candidate_careers/<int:candidate_id>/<int:page_no>', methods=['GET', 'POST'])
def edit_candidate_careers(candidate_id, page_no):
    if 'user_id' in session and 'user_type' in session:
        user_id = session['user_id']
        user_type = session['user_type']
        user_name = session['user_name']
        count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                          Notification.recruiter_name == user_name).count()
        career_count_notification_no = Career_notification.query.filter(
            Career_notification.notification_status == 'false',
            Career_notification.recruiter_name == user_name).count()

        if request.method == 'POST':
            # Retrieve the logged-in user's ID and user type from the session
            user_id = session['user_id']
            user_type = session['user_type']

            # Retrieve the form data for the candidate
            candidate = Candidate.query.get(candidate_id)

            # Update the candidate information based on user type
            if user_type == 'recruiter':
                candidate.recruiter = User.query.get(user_id).name
            elif user_type == 'management':
                candidate.management = User.query.get(user_id).name

            # Update the candidate fields with the new form data
            candidate.name = request.form.get('name')
            candidate.mobile = request.form.get('mobile')
            candidate.email = request.form.get('email')
            candidate.client = request.form.get('client')
            candidate.current_company = request.form.get('current_company')
            candidate.position = request.form.get('position')
            candidate.profile = request.form.get('profile')
            candidate.current_job_location = request.form.get('current_job_location')
            candidate.preferred_job_location = request.form.get('preferred_job_location')
            candidate.qualifications = request.form.get('qualifications')
            experience = request.form.get('experience')
            exp_months = request.form.get('exp_months')
            candidate.experience = experience +'.'+exp_months
            relevant_experience = request.form.get('relevant_experience')
            relevant_exp_months = request.form.get('relevant_exp_months')
            candidate.relevant_experience = relevant_experience + '.' + relevant_exp_months
            candidate.current_ctc = request.form.get('current_ctc')
            candidate.expected_ctc = request.form.get('expected_ctc')
            currency_type_current = request.form['currency_type_current']
            currency_type_except = request.form['currency_type_except']
            candidate.current_ctc = currency_type_current + " " + request.form['current_ctc']
            candidate.expected_ctc = currency_type_except + " " + request.form['expected_ctc']
            candidate.notice_period = request.form.get('notice_period')
            candidate.reason_for_job_change = request.form.get('reason_for_job_change')
            candidate.linkedin_url = request.form.get('linkedin')
            candidate.remarks = request.form.get('remarks')
            candidate.skills = request.form.get('skills')
            candidate.holding_offer = request.form.get('holding_offer')
            candidate.total = request.form.get('total')
            candidate.package_in_lpa = request.form.get('package_in_lpa')
            candidate.period_of_notice = request.form.get('period_of_notice')

            # Handle the resume file upload
            resume_file = request.files['resume']
            if resume_file.filename != '':
                # Save the new resume to the candidate's resume field as bytes
                candidate.resume = resume_file.read()

            holding_offer = request.form.get('holding_offer')
            if holding_offer == 'yes':
                total = request.form.get('total')
                package_in_lpa = request.form.get('package_in_lpa')

                candidate.total = total
                candidate.package_in_lpa = package_in_lpa
            elif holding_offer in ['no', 'pipeline']:
                candidate.total = None
                candidate.package_in_lpa = None

            notice_period = request.form.get('notice_period')
            if notice_period == 'yes':
                last_working_date = request.form['last_working_date']
                buyout = 'buyout' in request.form
                candidate.last_working_date = last_working_date
                candidate.buyout = buyout
            elif notice_period == 'no':
                period_of_notice = request.form['months']
                buyout = 'buyout' in request.form
                candidate.period_of_notice = period_of_notice
                candidate.buyout = buyout
            elif notice_period == 'completed':
                last_working_date = request.form['last_working_date']
                candidate.last_working_date = last_working_date

            db.session.commit()

            return redirect(
                url_for('career_dashboard', page_no=page_no, edit_candidate_message='Candidate Details Edited Successfully'))

        candidate = Candidate.query.get(candidate_id)
        candidate_data = {
            'id': candidate.id,
            'name': candidate.name,
            'mobile': candidate.mobile,
            'email': candidate.email,
            'client': candidate.client,
            'current_company': candidate.current_company,
            'position': candidate.position,
            'profile': candidate.profile,
            'current_job_location': candidate.current_job_location,
            'preferred_job_location': candidate.preferred_job_location,
            'qualifications': candidate.qualifications,
            'experience': candidate.experience,
            'relevant_experience': candidate.relevant_experience,
            'current_ctc': candidate.current_ctc,
            'expected_ctc': candidate.expected_ctc,
            'notice_period': candidate.notice_period,
            'reason_for_job_change': candidate.reason_for_job_change,
            'remarks': candidate.remarks,
            'candidate_status': candidate.status,
            'linkedin_url': candidate.linkedin_url,
            'skills': candidate.skills,
            'resume': candidate.resume,
            'holding_offer': candidate.holding_offer,
            'total': candidate.total,
            'package_in_lpa': candidate.package_in_lpa,
            'last_working_date': candidate.last_working_date,
            'buyout': candidate.buyout,
            'period_of_notice': candidate.period_of_notice,
        }

        return render_template('edit_candidate_careers.html', candidate_data=candidate_data, user_id=user_id,
                               user_type=user_type, user_name=user_name, count_notification_no=count_notification_no,
                               page_no=page_no,career_count_notification_no=career_count_notification_no)

    return redirect(url_for('career_dashboard'))


# Function to check if a filename has an allowed extension
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


from flask import send_file

import io
import os
import base64
from flask import redirect, url_for, send_file

@app.route('/download_resume/<int:candidate_id>', methods=['GET','POST'])
def download_resume(candidate_id):
    candidate = Candidate.query.get(candidate_id)

    if candidate is None or candidate.resume is None:
        return redirect(url_for('dashboard'))

    # Decode the base64 encoded resume
    resume_data = candidate.resume.split(',')[1]  # Get the data part after the comma
    resume_bytes = base64.b64decode(resume_data)

    # Determine the file extension
    if candidate.resume.startswith("data:application/pdf"):
        resume_filename = f"{candidate.name}_resume.pdf"
    else:
        resume_filename = f"{candidate.name}_resume.docx"

    # Send the resume data for download
    return send_file(io.BytesIO(resume_bytes),
                     attachment_filename=resume_filename,
                     as_attachment=True)


@app.route('/post_job', methods=['POST'])
def post_job():
    try:
        # Accessing the JSON data from the request
        data = request.json
        user_id = data['user_id']
        user = User.query.filter_by(id=user_id).first()
        print("user :",user)
        user_name = user.username
        print("user_name:",user_name)
        # Check if the "user_name" field exists
        if user_name:
            user_type = user.user_type

            if user_type == 'management':
                client = data.get('client')
                experience_min = data.get('experience_min')
                experience_max = data.get('experience_max')
                budget_min = data.get('budget_min')
                budget_max = data.get('budget_max')
                currency_type_min = data.get('currency_type_min')
                currency_type_max = data.get('currency_type_max')
                budget_min = currency_type_min + ' ' + budget_min
                budget_max = currency_type_max + ' ' + budget_max
                location = data.get('location')
                shift_timings = data.get('shift_timings')
                notice_period = data.get('notice_period')
                role = data.get('role')
                detailed_jd = data.get('detailed_jd')
                mode = data.get('mode')
                job_status = data.get('job_status')
                job_type = data.get('job_type')
                skills = data.get('skills')
                jd_pdf = data.get('jd_pdf')
                # Job_Type_details=data.get('Job_Type_details')

                if job_type == 'Contract':
                    Job_Type_details = data.get('Job_Type_details')
                    job_type = job_type + '(' + Job_Type_details + ' Months )'

                recruiter_names = data.get('recruiter', [])
                joined_recruiters = ', '.join(recruiter_names)

                new_job_post = JobPost(
                    client=client,
                    experience_min=experience_min,
                    experience_max=experience_max,
                    budget_min=budget_min,
                    budget_max=budget_max,
                    location=location,
                    shift_timings=shift_timings,
                    notice_period=notice_period,
                    role=role,
                    detailed_jd=detailed_jd,
                    mode=mode,
                    recruiter=joined_recruiters,
                    management=user.username,
                    job_status=job_status,
                    job_type=job_type,
                    skills=skills,
                    jd_pdf=jd_pdf
                )

                new_job_post.notification = 'no'
                new_job_post.date_created = date.today()
                new_job_post.time_created = datetime.now().time()

                # Add the new_job_post to the session and commit to generate the job_post_id
                db.session.add(new_job_post)
                db.session.commit()

                # Generate job_post_id after committing the new_job_post
                job_post_id = new_job_post.id

                # Define an empty list to hold Notification instances
                notifications = []

                for recruiter_name in joined_recruiters.split(','):
                    notification_status = False
                    notification = Notification(
                        job_post_id=job_post_id,  # Add job_post_id to Notification
                        recruiter_name=recruiter_name.strip(),
                        notification_status=notification_status
                    )
                    # Append each Notification instance to the notifications list
                    notifications.append(notification)

                # Add the notifications to the session and commit
                db.session.add_all(notifications)
                db.session.commit()
                notifications = Notification.query.filter_by(job_post_id=job_post_id).all()
                for notification in notifications:
                    notification.num_notification += 1
                db.session.commit()

                # Retrieve the email addresses of the recruiters
                recruiter_emails = [recruiter.email for recruiter in User.query.filter(User.username.in_(recruiter_names),
                                                                                         User.user_type == 'recruiter',
                                                                                         User.is_active == True,
                                                                                         User.is_verified == True)]
                for email in recruiter_emails:
                    send_notification(email)

                # Return the job_id along with the success message
                return jsonify({"message": "Job posted successfully", "job_id": job_post_id}), 200
            else:
                return jsonify({"error": "Invalid user type"}), 400
        else:
            return jsonify({"error": "Missing 'user_name' field in the request"}), 400

    except KeyError as e:
        return jsonify({"error": f"KeyError: {e}"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# @app.route('/post_job', methods=['POST'])
# def post_job():
#     try:
#         # Accessing the JSON data from the request
#         data = request.json
#         user_id = data['user_id']
#         user = User.query.filter_by(id=user_id).first()
#         user_name = user.username
#         # Check if the "user_name" field exists
#         if user_name:
#             user_type = user.user_type

#             if user_type == 'management':
#                 client = data.get('client')
#                 experience_min = data.get('experience_min')
#                 experience_max = data.get('experience_max')
#                 budget_min = data.get('budget_min')
#                 budget_max = data.get('budget_max')
#                 currency_type_min = data.get('currency_type_min')
#                 currency_type_max = data.get('currency_type_max')
#                 budget_min = currency_type_min + ' ' + budget_min if currency_type_min and budget_min else budget_min
#                 budget_max = currency_type_max + ' ' + budget_max if currency_type_max and budget_max else budget_max
#                 location = data.get('location')
#                 shift_timings = data.get('shift_timings')
#                 notice_period = data.get('notice_period')
#                 role = data.get('role')
#                 detailed_jd = data.get('detailed_jd')
#                 mode = data.get('mode')
#                 job_status = data.get('job_status')
#                 job_type = data.get('job_type')
#                 skills = data.get('skills')
#                 jd_pdf = data.get('jd_pdf')

#                 # Debug print statements
#                 print(f"Received job_type: {job_type}")
#                 print(f"Received data: {data}")

#                 if job_type == 'Contract':
#                     Job_Type_details = data.get('Job_Type_details')
#                     job_type = f'{job_type} ({Job_Type_details} Months)' if Job_Type_details else job_type

#                 recruiter_names = data.get('recruiter', [])
#                 joined_recruiters = ', '.join(recruiter_names)

#                 new_job_post = JobPost(
#                     client=client,
#                     experience_min=experience_min,
#                     experience_max=experience_max,
#                     budget_min=budget_min,
#                     budget_max=budget_max,
#                     location=location,
#                     shift_timings=shift_timings,
#                     notice_period=notice_period,
#                     role=role,
#                     detailed_jd=detailed_jd,
#                     mode=mode,
#                     recruiter=joined_recruiters,
#                     management=user.username,
#                     job_status=job_status,
#                     job_type=job_type,
#                     skills=skills,
#                     jd_pdf=jd_pdf
#                 )

#                 new_job_post.notification = 'no'
#                 new_job_post.date_created = date.today()
#                 new_job_post.time_created = datetime.now().time()

#                 # Define an empty list to hold Notification instances
#                 notifications = []

#                 if ',' in joined_recruiters:
#                     recruiter_names_lst = joined_recruiters.split(',')
#                     for recruiter_name in recruiter_names_lst:
#                         notification_status = False
#                         notification = Notification(
#                             recruiter_name=recruiter_name.strip(),
#                             notification_status=notification_status
#                         )
#                         # Append each Notification instance to the notifications list
#                         notifications.append(notification)
#                 else:
#                     recruiter_name = joined_recruiters
#                     notification_status = False
#                     notification = Notification(
#                         recruiter_name=recruiter_name,
#                         notification_status=notification_status
#                     )
#                     # Append each Notification instance to the notifications list
#                     notifications.append(notification)

#                 # Add the new_job_post and all associated notifications to the session
#                 db.session.add(new_job_post)
#                 db.session.add_all(notifications)
#                 db.session.commit()

#                 # Retrieve the email addresses of the recruiters
#                 recruiter_emails = [recruiter.email for recruiter in User.query.filter(User.username.in_(recruiter_names),
#                                                                                        User.user_type == 'recruiter',
#                                                                                        User.is_active == True,
#                                                                                        User.is_verified == True)]
#                 for email in recruiter_emails:
#                     send_notification(email)

#                 # Return the job_id along with the success message
#                 return jsonify({"message": "Job posted successfully", "job_id": new_job_post.id}), 200
#             else:
#                 return jsonify({"error": "Invalid user type"}), 400
#         else:
#             return jsonify({"error": "Missing 'user_name' field in the request"}), 400

#     except KeyError as e:
#         return jsonify({"error": f"KeyError: {e}"}), 400

#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


# @app.route('/recruiter_job_posts/<int:user_id>', methods=['GET'])
# def recruiter_job_posts(user_id):
    
#     if not user_id:
#         return jsonify({"error": "Missing user_id parameter"})

#     # Validate user existence
#     recruiter = User.query.get(user_id)
#     if not recruiter:
#         return jsonify({"error": "Recruiter not found"})

#     recruiter_name = recruiter.name

#     # Filter unread notifications efficiently using the recruiter's ID
#     unread_notifications = Notification.query.filter(
#         Notification.id == recruiter.id,
#         Notification.notification_status == False
#     ).all()

#     # Filter active and on-hold job posts
#     active_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter,
#         JobPost.job_status == 'Active'
#     ).order_by(JobPost.id).all()

#     on_hold_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter,
#         JobPost.job_status == 'Hold'
#     ).order_by(JobPost.id).all()

#     # Update notification statuses after retrieving them
#     for notification in unread_notifications:
#         notification.notification_status = True
#     db.session.commit()

#     # Construct JSON response with relevant data
#     response_data = {
#         "count_notification_no": len(unread_notifications),
#         "user_name": recruiter_name,
#         "job_posts": [
#             {  # Include only necessary job post fields
#                 "id": job_post.id,
#                 "title": job_post.title,
#                 "description": job_post.description,
#                 "created_at": job_post.created_at.isoformat()  # Example for date formatting
#             }
#             for job_post in active_job_posts
#         ],
#         "job_posts_hold": [
#             {  # Include only necessary job post fields (optional)
#                 "id": job_post.id,
#                 "title": job_post.title,
#                 "description": job_post.description,
#                 "created_at": job_post.created_at.isoformat()  # Example for date formatting
#             }
#             for job_post in on_hold_job_posts
#         ],
#         "career_count_notification_no": 0  # Placeholder, implement career notification logic
#     }

#     # # Include optional parameters conditionally
#     # if url_for('add_candidate'):
#     #     response_data["redirect_url"] = url_for('add_candidate')
#     # if request.args.get('no_doc_message'):
#     #     response_data["no_doc_message"] = request.args.get('no_doc_message')

#     return jsonify(response_data)


# @app.route('/recruiter_job_posts/<int:user_id>', methods=['GET'])
# def recruiter_job_posts(user_id):
   
#     if not user_id:
#         return jsonify({"error": "Missing user_id parameter"})

#     # Validate user existence
#     recruiter = User.query.get(user_id)
#     if not recruiter:
#         return jsonify({"error": "Recruiter not found"})

#     recruiter_name = recruiter.name

#     # Filter unread notifications efficiently using the recruiter's ID
#     unread_notifications = Notification.query.filter(
#         Notification.recruiter_id == recruiter.id,
#         Notification.notification_status == False
#     ).all()

#     # Filter active and on-hold job posts
#     active_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter,
#         JobPost.job_status == 'Active'
#     ).order_by(JobPost.id).all()

#     on_hold_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter,
#         JobPost.job_status == 'Hold'
#     ).order_by(JobPost.id).all()

#     # Update notification statuses after retrieving them
#     for notification in unread_notifications:
#         notification.notification_status = True
#     db.session.commit()

#     # Construct JSON response with serialized job post data
#     response_data = {
#         "count_notification_no": len(unread_notifications),
#         "job_posts": [job_post.serialize() for job_post in active_job_posts],
#         "user_name": recruiter_name,
#         "job_posts_hold": [job_post.serialize() for job_post in on_hold_job_posts],
#         "redirect_url": url_for('add_candidate'),  # Optional, include if needed
#         "no_doc_message": request.args.get('no_doc_message'),  # Optional, include if needed
#         "career_count_notification_no": 0  # Placeholder, implement career notification logic
#     }

#     return jsonify(response_data)

# @app.route('/recruiter_job_posts', methods=['POST'])
# def recruiter_job_posts():
#     data = request.json
#     user_id = data.get('user_id')  # Using get() to avoid KeyError if 'user_id' is missing
#     if not user_id:
#         return jsonify({"error": "User ID is missing"}), 400

#     # Validate user existence
#     recruiter = User.query.get(user_id)
#     if not recruiter:
#         return jsonify({"error": "Recruiter not found"}), 404

#     recruiter_name = recruiter.name

#     # Filter unread notifications based on recruiter name
#     unread_notifications = Career_notification.query.filter(
#         Career_notification.recruiter_name == recruiter_name,
#         Career_notification.notification_status == False
#     ).all()

#     # Filter active and on-hold job posts
#     active_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
#         JobPost.job_status == 'Active'
#     ).order_by(JobPost.id).all()

#     on_hold_job_posts = JobPost.query.filter(
#         JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
#         JobPost.job_status == 'Hold'
#     ).order_by(JobPost.id).all()

#     # Update notification statuses after retrieving them
#     for notification in unread_notifications:
#         notification.notification_status = True
#     db.session.commit()

#     # Construct JSON response
#     response_data = {
#         "count_notification_no": len(unread_notifications),
#         "job_posts": [job_post_to_dict(job_post) for job_post in active_job_posts],
#         "user_name": recruiter_name,
#         "job_posts_hold": [job_post_to_dict(job_post) for job_post in on_hold_job_posts],
#         "redirect_url": url_for('add_candidate'),  # Optional, include if needed
#         "no_doc_message": request.args.get('no_doc_message'),  # Optional, include if needed
#         "career_count_notification_no": 0  # Placeholder, implement career notification logic
#     }

#     return jsonify(response_data)

@app.route('/recruiter_job_posts', methods=['POST'])
def recruiter_job_posts():
    data = request.json
    user_id = data.get('user_id')  # Using get() to avoid KeyError if 'user_id' is missing
    if not user_id:
        return jsonify({"error": "User ID is missing"}), 400

    # Validate user existence
    recruiter = User.query.get(user_id)
    if not recruiter:
        return jsonify({"error": "Recruiter not found"}), 404

    recruiter_name = recruiter.name

    # Filter unread notifications based on recruiter name
    unread_notifications = Career_notification.query.filter(
        Career_notification.recruiter_name == recruiter_name,
        Career_notification.notification_status == False
    ).all()

    # Filter active and on-hold job posts
    active_job_posts = JobPost.query.filter(
        JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
        JobPost.job_status == 'Active'
    ).order_by(JobPost.id).all()

    on_hold_job_posts = JobPost.query.filter(
        JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
        JobPost.job_status == 'Hold'
    ).order_by(JobPost.id).all()

    # Update notification statuses after retrieving them
    for notification in unread_notifications:
        notification.notification_status = True
    db.session.commit()

    # Construct JSON response
    response_data = {
        "count_notification_no": len(unread_notifications),
        "job_posts": [job_post_to_dict(job_post) for job_post in active_job_posts],
        "user_name": recruiter_name,
        "job_posts_hold": [job_post_to_dict(job_post) for job_post in on_hold_job_posts],
        "redirect_url": url_for('add_candidate'),  # Optional, include if needed
        "no_doc_message": request.args.get('no_doc_message'),  # Optional, include if needed
        "career_count_notification_no": 0  # Placeholder, implement career notification logic
    }

    return jsonify(response_data)

# Helper function to convert JobPost object to dictionary
def job_post_to_dict(job_post):
    data_updated_date_str = job_post.data_updated_date.strftime('%Y-%m-%d') if job_post.data_updated_date else None
    data_updated_time_str = job_post.data_updated_time.strftime('%H:%M:%S') if job_post.data_updated_time else None

    return {
        "id": job_post.id,
        "client": job_post.client,
        "experience_min": job_post.experience_min,
        "experience_max": job_post.experience_max,
        "budget_min": job_post.budget_min,
        "budget_max": job_post.budget_max,
        "location": job_post.location,
        "shift_timings": job_post.shift_timings,
        "notice_period": job_post.notice_period,
        "role": job_post.role,
        "detailed_jd": job_post.detailed_jd,
        "mode": job_post.mode,
        "recruiter": job_post.recruiter,
        "management": job_post.management,
        "date_created": job_post.date_created.strftime('%Y-%m-%d'),
        "time_created": job_post.time_created.strftime('%H:%M:%S'),
        "job_status": job_post.job_status,
        "job_type": job_post.job_type,
        "skills": job_post.skills,
        "notification": job_post.notification,
        "data_updated_date": data_updated_date_str,
        "data_updated_time": data_updated_time_str
    }


from flask import jsonify

@app.route('/update_job_status/<int:job_id>', methods=['POST'])
def update_job_status(job_id):
    data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({"success": False, "error": "User not found"}), 404

    user_type = user.user_type
    username = user.username

    # Retrieve the job post from the database based on the provided job_id
    job_post = JobPost.query.get(job_id)

    if job_post:
        try:
            # Extract the new job status from the form data
            new_job_status = data['new_job_status']

            # Update the job status
            job_post.job_status = new_job_status

            # Commit the changes to the database
            db.session.commit()

            # Return a JSON response indicating success
            return jsonify({"success": True, "message": "Job status updated successfully"})
        
        except KeyError:
            # If 'new_job_status' key is missing in form data
            return jsonify({"success": False, "error": "Missing 'new_job_status' in form data"}), 400
        
        except Exception as e:
            # Handle other exceptions
            db.session.rollback()  # Rollback any changes made to the session
            return jsonify({"success": False, "error": str(e)}), 500

    # If job_post is None (job not found)
    return jsonify({"success": False, "error": "Job post not found"}), 404




import base64

@app.route('/view_all_jobs', methods=['POST'])
def view_all_jobs():
    # Get data from JSON request
    data = request.json

    # Extract any parameters you need from the JSON data
    user_name = data['username']

    # Retrieve all job posts from the database
    job_posts_active = JobPost.query.filter_by(job_status='Active').order_by(JobPost.id).all()
    job_posts_hold = JobPost.query.filter_by(job_status='Hold').order_by(JobPost.id).all()

    # Construct JSON response
    response_data = {
        "user_name": user_name,
        "job_posts_active": [
            {
                "id": job_post.id,
                "client": job_post.client,
                "role": job_post.role,
                "experience_min": job_post.experience_min,
                "experience_max": job_post.experience_max,
                "budget_min": job_post.budget_min,
                "budget_max": job_post.budget_max,
                "location": job_post.location,
                "shift_timings": job_post.shift_timings,
                "notice_period": job_post.notice_period,
                "detailed_jd": job_post.detailed_jd,
                "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
                "mode": job_post.mode,
                "recruiter": job_post.recruiter,
                "management": job_post.management,
                "job_status": job_post.job_status,
                "job_type": job_post.job_type,
                "skills": job_post.skills,
                "date_created": str(job_post.date_created),
                "time_created": str(job_post.time_created),
                "data_updated_date":str(job_post.data_updated_date),
                "data_updated_time":str(job_post.data_updated_time)
                # Include other attributes as needed
            }
            for job_post in job_posts_active
        ],
        "job_posts_hold": [
            {
                "id": job_post.id,
                "client": job_post.client,
                "role": job_post.role,
                "experience_min": job_post.experience_min,
                "experience_max": job_post.experience_max,
                "budget_min": job_post.budget_min,
                "budget_max": job_post.budget_max,
                "location": job_post.location,
                "shift_timings": job_post.shift_timings,
                "notice_period": job_post.notice_period,
                "detailed_jd": job_post.detailed_jd,
                "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
                "mode": job_post.mode,
                "recruiter": job_post.recruiter,
                "management": job_post.management,
                "job_status": job_post.job_status,
                "job_type": job_post.job_type,
                "skills": job_post.skills,
                "date_created": str(job_post.date_created),
                "time_created": str(job_post.time_created),
                "data_updated_date":str(job_post.data_updated_date),
                "data_updated_time":str(job_post.data_updated_time)
                # Include other attributes as needed
            }
            for job_post in job_posts_hold
        ]
    }

    # Return JSON response
    return jsonify(response_data)


# @app.route('/view_all_jobs', methods=['POST'])
# def view_all_jobs():
#     # Get data from JSON request
#     data = request.json

#     # Extract any parameters you need from the JSON data
#     user_name = data['username']

#     # Retrieve all job posts from the database
#     job_posts_active = JobPost.query.filter_by(job_status='Active').order_by(JobPost.id).all()
#     job_posts_hold = JobPost.query.filter_by(job_status='Hold').order_by(JobPost.id).all()

#     # Construct JSON response
#     response_data = {
#         "user_name": user_name,
#         "job_posts_active": [
#             {
#                 "id": job_post.id,
#                 "client": job_post.client,
#                 "role": job_post.role,
#                 "experience_min": job_post.experience_min,
#                 "experience_max": job_post.experience_max,
#                 "budget_min": job_post.budget_min,
#                 "budget_max": job_post.budget_max,
#                 "location": job_post.location,
#                 "shift_timings": job_post.shift_timings,
#                 "notice_period": job_post.notice_period,
#                 "detailed_jd": job_post.detailed_jd,
#                 "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
#                 "mode": job_post.mode,
#                 "recruiter": job_post.recruiter,
#                 "management": job_post.management,
#                 "job_status": job_post.job_status,
#                 "job_type": job_post.job_type,
#                 "skills": job_post.skills,
#                 "date_created": str(job_post.date_created),
#                 "time_created": str(job_post.time_created),
#                 "data_updated_date":str(job_post.data_updated_date),
#                 "data_updated_time":str(jobs_post.data_updated_time)
#                 # Include other attributes as needed
#             }
#             for job_post in job_posts_active
#         ],
#         "job_posts_hold": [
#             {
#                 "id": job_post.id,
#                 "client": job_post.client,
#                 "role": job_post.role,
#                 "experience_min": job_post.experience_min,
#                 "experience_max": job_post.experience_max,
#                 "budget_min": job_post.budget_min,
#                 "budget_max": job_post.budget_max,
#                 "location": job_post.location,
#                 "shift_timings": job_post.shift_timings,
#                 "notice_period": job_post.notice_period,
#                 "detailed_jd": job_post.detailed_jd,
#                 "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
#                 "mode": job_post.mode,
#                 "recruiter": job_post.recruiter,
#                 "management": job_post.management,
#                 "job_status": job_post.job_status,
#                 "job_type": job_post.job_type,
#                 "skills": job_post.skills,
#                 "date_created": str(job_post.date_created),
#                 "time_created": str(job_post.time_created),
#                 "data_updated_date":str(job_post.data_updated_date),
#                 "data_updated_time":str(jobs_post.data_updated_time)
#                 # Include other attributes as needed
#             }
#             for job_post in job_posts_hold
#         ]
#     }

#     # Return JSON response
#     return jsonify(response_data)

def send_notification(recruiter_email):
    msg = Message('New Job Posted', sender='ganesh.s@makonissoft.com', recipients=[recruiter_email])
    msg.body = 'A new job has been posted. Check your dashboard for more details.'
    mail.send(msg)

@app.route('/other_job_posts', methods=['GET'])
def other_job_posts():
    if 'user_id' in session and 'user_type' in session:
        if session['user_type'] == 'recruiter':
            # Retrieve the logged-in user's ID from the session
            user_id = session['user_id']

            # Retrieve the recruiter's name based on user ID
            recruiter_name = User.query.get(user_id).name

            job_posts = JobPost.query.filter(JobPost.recruiter != recruiter_name).distinct(JobPost.client).all()

            return render_template('other_job_posts.html', job_posts=job_posts)

    # Redirect or render an appropriate page if the conditions are not met
    return redirect(url_for('login'))

# @app.route('/recruiter_job_posts', methods=['GET'])
# def recruiter_job_posts():
#     no_doc_message = request.args.get('no_doc_message')
#     if 'user_id' in session and 'user_type' in session:
#         if session['user_type'] == 'recruiter':
#             # Retrieve the logged-in user's ID from the session
#             user_id = session['user_id']
#             user_name = session['user_name']
#             count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
#                                                               Notification.recruiter_name == user_name).count()
#             career_count_notification_no = Career_notification.query.filter(Career_notification.notification_status == 'false',
#                                                               Career_notification.recruiter_name == user_name).count()
#             recruiter_name = User.query.get(user_id).name

#             job_posts = JobPost.query.filter(JobPost.recruiter.contains(recruiter_name),
#                                              JobPost.job_status == 'Active').order_by(JobPost.id).all()
#             job_posts_hold = JobPost.query.filter(JobPost.recruiter.contains(recruiter_name),
#                                                   JobPost.job_status == 'Hold').order_by(JobPost.id).all()

#             notifications = Notification.query.filter(Notification.recruiter_name.contains(recruiter_name)).all()

#             for notification in notifications:
#                 if notification.notification_status == False:
#                     notification.notification_status = True
#                     db.session.commit()

#             # for job_post in job_posts:
#             #     if job_post.notification == 'no':
#             #         job_post.notification = 'yes'
#             #         db.session.commit()

#             return render_template('recruiter_job_posts.html', count_notification_no=count_notification_no,
#                                    job_posts=job_posts, user_name=user_name, job_posts_hold=job_posts_hold,
#                                    redirect_url=url_for('add_candidate'), recruiter_job_posts=recruiter_job_posts,
#                                    no_doc_message=no_doc_message, career_count_notification_no=career_count_notification_no)

#     return redirect(url_for('login'))
    

import base64
import io
import magic

@app.route('/view_resume/<int:candidate_id>', methods=['GET'])
def view_resume(candidate_id):
    # Retrieve the resume data from the database using SQLAlchemy
    candidate = Candidate.query.filter_by(id=candidate_id).first()
    if not candidate:
        return 'Candidate not found'
    # Decode the base64 encoded resume data
    print("candidate.resume",candidate.resume.tobytes())
    if "==" not in str(candidate.resume.tobytes()):
        if request.args.get('decode') == 'base64':
            # Decode the base64 encoded resume data
            decoded_resume = base64.b64decode(candidate.resume)
            resume_binary = decoded_resume
        else:
            # Retrieve the resume binary data from the database
            resume_binary = candidate.resume.tobytes()  # Convert memoryview to bytes

        # Determine the mimetype based on the file content
        is_pdf = resume_binary.startswith(b"%PDF")
        mimetype = 'application/pdf' if is_pdf else 'application/msword'

        # Send the file as a response
        return send_file(
            io.BytesIO(resume_binary),
            mimetype=mimetype,
            as_attachment=False
        )
    else:
        decoded_resume = base64.b64decode(candidate.resume)
        # Create a file-like object (BytesIO) from the decoded resume data
        resume_file = io.BytesIO(decoded_resume)
        # Determine the mimetype based on the file content
        is_pdf = decoded_resume.startswith(b"%PDF")
        mimetype = 'application/pdf' if is_pdf else 'application/msword'

        # Send the file as a response
        return send_file(
            resume_file,
            mimetype=mimetype,
            as_attachment=False
        )



# @app.route('/view_resume/<int:candidate_id>', methods=['GET'])
# def view_resume(candidate_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return 'Candidate not found', 404

#     # Check if the request specifies to retrieve the resume data directly from the database
#     if request.args.get('decode') == 'base64':
#         # Decode the base64 encoded resume data
#         decoded_resume = base64.b64decode(candidate.resume)
#         resume_binary = decoded_resume
#     else:
#         # Retrieve the resume binary data from the database
#         resume_binary = candidate.resume.tobytes()  # Convert memoryview to bytes

#     # Determine the mimetype based on the file content
#     is_pdf = resume_binary.startswith(b"%PDF")
#     mimetype = 'application/pdf' if is_pdf else 'application/msword'

#     # Send the file as a response
#     return send_file(
#         io.BytesIO(resume_binary),
#         mimetype=mimetype,
#         as_attachment=False
#     )


 
# @app.route('/view_resume/<int:candidate_id>', methods=['GET'])
# def view_resume(candidate_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return 'Candidate not found'
#     # Decode the base64 encoded resume data
        
#     decoded_resume = base64.b64decode(candidate.resume)
#     # Create a file-like object (BytesIO) from the decoded resume data
#     resume_file = io.BytesIO(decoded_resume)
#     # Determine the mimetype based on the file content
#     is_pdf = decoded_resume.startswith(b"%PDF")
#     mimetype = 'application/pdf' if is_pdf else 'application/msword'
 
#     # Send the file as a response
#     return send_file(
#         resume_file,
#         mimetype=mimetype,
#         as_attachment=False
#     )
###############################################################################################
# @app.route('/upload_user_image/<int:user_id>', methods=['POST'])
# def upload_user_image(user_id):
#     data=request.json
#     # if not data or 'image' not in data:
#     #     return jsonify({'error': 'No image data provided'}), 400
    
#     filename=data['file_name']
#     image_file=data['image_file']

#     user = User.query.filter_by(id=user_id).first()
#     # if not user:
#     #     return jsonify({'error': 'User not found'}), 400
    
#     user.filename = filename
#     user.image_file=image_file
#     db.session.commit()

#     return jsonify({'message': 'Image updated successfully'}), 200
#######################################################################################


# @app.route('/upload_user_image/<int:user_id>', methods=['POST'])
# def upload_user_image(user_id):
#     data = request.json
#     if not data:
#         return jsonify({'error': 'Invalid JSON data provided'}), 400

#     image_content = data.get('image')
#     file_name = data.get('filename')

#     # if not image_content or not file_name:
#     #     return jsonify({'error': 'Image content or filename missing in the request'}), 400

#     # Find the user by user_id
#     user = User.query.get(user_id)
#     if not user:
#         return jsonify({'error': 'User not found'}), 404

#     # Update user's filename and image content
#     user.filename = file_name
#     user.image_file = image_content

#     # Commit changes to the database
#     db.session.commit()

#     return jsonify({'message': 'Image updated successfully'}), 200

# @app.route('/upload_user_image/<int:user_id>', methods=['POST'])
# def upload_user_image(user_id):
#     data = request.json
#     print("\n\n\n\n\n")
#     print("Data :",data)
#     image_content = data['image']
#     file_name = data['filename']
#     # Find the user by user_id
#     user = User.query.get(user_id)
#     if not user:
#         return jsonify({'error': 'User not found'}), 404

#     # Update user's filename and image content
#     user.filename = file_name
#     user.image_file = image_content
#     # db.session.commit()

#     # Commit changes to the database
#     try:
#         db.session.commit()
#     except:
#         print("Failed to Upload !!")

#     return jsonify({'message': 'Image updated successfully'}), 200


# @app.route('/upload_user_image/<int:user_id>', methods=['POST'])
# def upload_user_image(user_id):
#     data = request.form

#     # Extract file name and image content
#     image_content = data['image']
#     filename = data['filename']  # Retrieve file object
     
#     # Find the user by user_id
#     user = User.query.get(user_id)
#     if not user:
#         return jsonify({'error': 'User not found'}), 404

#     # Update user's filename and image content
#     user.filename = filename
#     user.image_file = image_content

#     # Commit changes to the database
#     db.session.commit()

#     return jsonify({'message': 'Image updated successfully'}), 200

@app.route('/upload_user_image/<int:user_id>', methods=['POST'])
def upload_user_image(user_id):
    try:
        # Extract file from request
        data=request.json
        image_file = data['image']
        filename = data['filename']
        image_delete_status=data['image_delete_status']
        
        # Find the user by user_id
        user = User.query.get(user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404

        # Update user's filename and image content
        user.filename = filename
        user.image_file = image_file  # Store image content as binary data
        user.image_deleted=image_delete_status
        # Commit changes to the database
        db.session.commit()

        return jsonify({'message': 'Image updated successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

#################################################################################################

import base64
import io
# @app.route('/image_status/<int:user_id>', methods=['GET'])
# def image_status(user_id):
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'error': 'Image not found'}), 404

#     return jsonify({'message': user.image_file}), 200


import io
import base64
import mimetypes

# @app.route('/user_image/<int:user_id>', methods=['GET'])
# def user_image(user_id):
#     # Retrieve the user data from the database
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'message': 'Image not found'}), 400
    
#     # Decode the bytea image data
#     image_data = user.image_file
#     image_file = base64.b64decode(image_data)
    
#     # Determine the MIME type dynamically
#     mime_type, _ = mimetypes.guess_type(user.image_filename)  # Assuming user.image_filename holds the filename
    
#     # Default to 'application/octet-stream' if MIME type couldn't be guessed
#     if not mime_type:
#         mime_type = 'application/octet-stream'
    
#     # Send the file as a response
#     return send_file(
#         io.BytesIO(image_file),
#         mimetype=mime_type,
#         as_attachment=False
#     )

import io
import base64
from PIL import Image
import mimetypes

@app.route('/user_image/<int:user_id>', methods=['GET'])
def user_image(user_id):
    # Retrieve the user data from the database
    user = User.query.filter_by(id=user_id).first()
    if not user or not user.image_file:
        return jsonify({'message': 'Image not found'}), 400
    
    # Decode the bytea image data
    image_data = base64.b64decode(user.image_file)
    
    # Determine the MIME type
    image = Image.open(io.BytesIO(image_data))
    mime_type = Image.MIME.get(image.format)
    
    # Send the file as a response
    return send_file(
        io.BytesIO(image_data),
        mimetype=mime_type,
        as_attachment=False
    )

# @app.route('/user_image/<int:user_id>', methods=['GET'])
# def user_image(user_id):
#     # Retrieve the user data from the database
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'message': 'Image not found'}),400
    
#     # Decode the bytea image data
#     image_data = user.image_file
    
#     image_file = base64.b64decode(image_data)
    
#     # Send the file as a response
#     return send_file(
#         io.BytesIO(image_file),
#          mimetype = 'image/jpeg',
#         as_attachment=False
#     )



# import base64
# import io
# @app.route('/user_image/<int:user_id>', methods=['GET'])
# def user_image(user_id):
#     # Retrieve the user data from the database
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'error': 'Image not found'}), 404
    
#     # Decode the bytea image data
#     image_data = user.image_file

#     # Create a file-like object (BytesIO) from the image data
#     image_file = io.BytesIO(image_data)

#     # Determine the mimetype based on the file content
#     if image_data.startswith(b"\x89PNG"):
#         mimetype = 'image/png'
#     elif image_data.startswith(b"\xff\xd8\xff"):
#         mimetype = 'image/jpeg'
#     elif image_data.startswith(b"\xff\xd8\xff\xe0") and image_data[6:10] in (b"JFIF", b"Exif"):
#         mimetype = 'image/jpeg'
#     else:
#         return jsonify({'error': 'Unsupported image format'}), 400

#     # Send the file as a response
#     return send_file(
#         image_file,
#         mimetype=mimetype,
#         as_attachment=False
#     )

@app.route('/delete_user_image/<int:user_id>', methods=['POST'])
def delete_user_image(user_id):
    data = request.json
    profile_image = data['profileImage']
    image_delete_status=data['image_delete_status']
    if not profile_image:
        return jsonify({"error": "Profile image must be specified"}), 400

    user = User.query.filter_by(id=user_id).first()
    
    if not user:
        return jsonify({"error": "User not found"}), 404

    user.image_file = None
    user.filename = None
    user.image_deleted = image_delete_status
    db.session.commit()

    return jsonify({"message": "Image file deleted successfully"}), 200


# @app.route('/delete_user_image/<int:user_id>', methods=['POST'])
# def delete_user_image(user_id):
#     data = request.json
#     image_file = data.get('image_file')
#     if not image_file:
#         return jsonify({"error": "Image file must be specified"}), 400

#     user = User.query.filter_by(id=user_id,).first()
    
#     if not user:
#         return jsonify({"error": "User not found"}), 400
    
#     # if user.image_file != image_file:
#     #     return jsonify({"error": "Image file does not match the user's image"}), 400

#     user.image_file = None
#     user.filename=None
#     db.session.commit()

#     return jsonify({"message": "Image file deleted successfully"}), 200
    

@app.route('/viewfull_jd/<int:id>')
def viewfull_jd(id):
    user_type = session['user_type']
    job_post = JobPost.query.get(id)
    return render_template('viewfull_jd.html', job_post=job_post,user_type=user_type)

@app.route('/add_candidate_view')
def add_candidate_view():
    user_id = session['user_id']
    user_type = session['user_type']
    user_name = session['user_name']

    if user_type == 'recruiter':
        recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
        if recruiter:
            candidates = Candidate.query.filter_by(
                recruiter=recruiter.name).all()  # Filter candidates by recruiter's name
            # data = json.dumps(candidates, sort_keys=False)
            results = db.session.query(JobPost.client, JobPost.recruiter).filter(
                JobPost.recruiter.contains(user_name)).all()
            client_names = sorted(list(set([result.client for result in results])))
            count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                              Notification.recruiter_name == user_name).count()
            return render_template('add_candidate_view.html', user=recruiter, user_type=user_type, user_name=user_name,
                                   candidates=candidates, count_notification_no=count_notification_no,
                                   client_names=client_names)
    elif user_type == 'management':
        users = User.query.all()
        candidates = Candidate.query.all()
        JobsPosted = JobPost.query.all()
        clients = db.session.query(JobPost.client).all()
        client_names = list(set([client[0] for client in clients]))

        return render_template('add_candidate_view.html', users=users, user_type=user_type, user_name=user_name,
                               JobsPosted=JobsPosted, client_names=client_names)

import os
import shutil
from flask import Flask, request, send_file, redirect, url_for
from zipfile import ZipFile

@app.route('/download_resumes')
def download_resumes():
    candidate_ids = request.args.getlist('candidate_ids')
    
    # Create a temporary directory to store resume files
    temp_dir = 'temp_resumes'
    os.makedirs(temp_dir, exist_ok=True)
    
    resume_paths = []

    for candidate_id in candidate_ids:
        candidate = Candidate.query.get(candidate_id)
        if candidate is None or candidate.resume is None:
            continue
        
        resume_file = io.BytesIO(candidate.resume)
        is_pdf = resume_file.getvalue().startswith(b"%PDF")
        if is_pdf : 
            resume_filename = f"{candidate.name}_resume.pdf" 
            resume_path = os.path.join(temp_dir, resume_filename)
            with open(resume_path, 'wb') as file:
                file.write(candidate.resume)
            
            resume_paths.append(resume_path)
        else:
            resume_filename = f"{candidate.name}_resume.docx" 
            resume_path = os.path.join(temp_dir, resume_filename)
            with open(resume_path, 'wb') as file:
                file.write(candidate.resume)
            
            resume_paths.append(resume_path)

    # Create a zip file containing all resume files
    zip_filename = 'resumes.zip'
    with ZipFile(zip_filename, 'w') as zipf:
        for resume_path in resume_paths:
            zipf.write(resume_path, os.path.basename(resume_path))
    
    # Clean up temporary directory
    shutil.rmtree(temp_dir)
    
    # Send the zip file for download
    return send_file(zip_filename, as_attachment=True)


@app.route('/assign_job/<int:job_id>', methods=['POST'])
def assign_job(job_id):
    data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({"error_message": "User not found"}), 404

    user_type = user.user_type
    username = user.username
    job_post = JobPost.query.get(job_id)  # Retrieve the job post by its ID

    if not job_post:
        return jsonify({"error_message": "Job not found"}), 404

    current_recruiters = job_post.recruiter.split(', ') if job_post.recruiter else []

    if request.method == 'POST':
        new_recruiter_names = data.get('recruiters', [])
        
        # Modification: Remove duplicate recruiters by combining lists and converting to a set
        updated_recruiter_names = list(set(current_recruiters + new_recruiter_names))
        
        # Join the recruiter names into a single string
        joined_recruiters = ', '.join(updated_recruiter_names)
        job_post.recruiter = joined_recruiters
        db.session.commit()

        # Send notification emails to the newly assigned recruiters
        new_recruiter_emails = [recruiter.email for recruiter in
                                User.query.filter(User.name.in_(new_recruiter_names),
                                                  User.user_type == 'recruiter')]
        for email in new_recruiter_emails:
            send_notification(email)

        # Define an empty list to hold Notification instances
        notifications = []

        for recruiter_name in updated_recruiter_names:
            if recruiter_name.strip() in new_recruiter_names:
                notification_status = False  # Set the initial status
                notification = Notification(
                    job_post_id=job_post.id,
                    recruiter_name=recruiter_name.strip(),
                    notification_status=notification_status
                )
                # Append each Notification instance to the notifications list
                notifications.append(notification)

        # Commit the notifications to the database session
        db.session.add_all(notifications)
        db.session.commit()

        return jsonify({"message": "Job re-assigned successfully"}), 200

    recruiter_names = [recruiter.name for recruiter in User.query.filter_by(user_type='recruiter')]
    return jsonify({
        "user_name": username,
        "job_post": job_post.serialize(),
        "current_recruiters": current_recruiters,
        "recruiters": recruiter_names
    })


# @app.route('/assign_job/<int:job_id>', methods=['POST'])
# def assign_job(job_id):
#     data = request.json
#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
    
#     if not user:
#         return jsonify({"error_message": "User not found"}), 404
    
#     user_type = user.user_type
#     username = user.username
#     job_post = JobPost.query.get(job_id)  # Retrieve the job post by its ID

#     if not job_post:
#         return jsonify({"error_message": "Job not found"}), 404

#     current_recruiters = job_post.recruiter.split(', ') if job_post.recruiter else []

#     if request.method == 'POST':
#         new_recruiter_names = data.get('recruiters', [])
#         all_recruiter_names = current_recruiters + new_recruiter_names
#         joined_recruiters = ', '.join(all_recruiter_names)
#         job_post.recruiter = joined_recruiters
#         db.session.commit()

#         # Send notification emails to the newly assigned recruiters
#         new_recruiter_emails = [recruiter.email for recruiter in
#                                 User.query.filter(User.name.in_(new_recruiter_names),
#                                                     User.user_type == 'recruiter')]
#         for email in new_recruiter_emails:
#             send_notification(email)

#         # Define an empty list to hold Notification instances
#         notifications = []

#         if ',' in joined_recruiters:
#             recruiter_names_lst = joined_recruiters.split(',')
#             for recruiter_name in recruiter_names_lst:
#                 if recruiter_name.strip() in new_recruiter_names:
#                     notification_status = False  # Set the initial status
#                     notification = Notification(
#                         recruiter_name=recruiter_name.strip(),
#                         notification_status=notification_status
#                     )
#                     # Append each Notification instance to the notifications list
#                     notifications.append(notification)
#         else:
#             recruiter_name = joined_recruiters
#             if recruiter_name in new_recruiter_names:
#                 notification_status = False  # Set the initial status
#                 notification = Notification(
#                     recruiter_name=recruiter_name,
#                     notification_status=notification_status
#                 )
#                 # Append each Notification instance to the notifications list
#                 notifications.append(notification)

#         # Commit the notifications to the database session
#         db.session.add_all(notifications)
#         db.session.commit()
#         return jsonify({"message": "Job re-assigned successfully"}), 200

#     recruiter_names = [recruiter.name for recruiter in User.query.filter_by(user_type='recruiter')]
#     return jsonify({
#         "user_name": username,
#         "job_post": job_post.serialize(),
#         "current_recruiters": current_recruiters,
#         "recruiters": recruiter_names
#     })



@app.route('/assign_candidate', methods=['POST'])
def assign_candidate():
    assignment_message = request.args.get('assignment_message')
    if 'user_id' in session and 'user_type' in session and session['user_type'] == 'management':
        user_name = session['user_name']
        recruiters = User.query.filter_by(user_type='recruiter').all()

        selected_recruiter_id = request.json.get('selected_recruiter_id')
        selected_candidate_id = request.json.get('selected_candidate_id')
        assign_recruiter_id = request.json.get('assign_recruiter_id')
        assign_candidate_id = request.json.get('assign_candidate_id')

        selected_recruiter = None
        selected_candidate = None
        assigned_recruiter = None

        if selected_recruiter_id:
            selected_recruiter = User.query.get(selected_recruiter_id)

        if selected_candidate_id:
            selected_candidate = Candidate.query.get(selected_candidate_id)

        if assign_recruiter_id:
            assigned_recruiter = User.query.get(assign_recruiter_id)

        if request.method == 'POST':
            selected_candidate_ids = request.json.get('selected_candidate_ids')
            if assigned_recruiter and selected_candidate_ids:
                for candidate_id in selected_candidate_ids:
                    candidate = Candidate.query.get(candidate_id)
                    if candidate:
                        candidate.recruiter = assigned_recruiter.name
                db.session.commit()
                return jsonify({"message": "Candidates Assigned Successfully"}), 200

            if selected_candidate_id:
                selected_candidate = Candidate.query.get(selected_candidate_id)

        candidates = []
        if selected_recruiter:
            candidates = Candidate.query.filter(
                Candidate.recruiter == selected_recruiter.name,
                Candidate.status.in_(['None', "SCREENING","L1 - SCHEDULED" ,"L1 - SELECTED", 'L1 - FEEDBACK', 'L1 - RESCHEDULE',"L2 - SCHEDULED" ,"L2 - SELECTED",
                                      'L2 - FEEDBACK', 'L2 - RESCHEDULE', 'HOLD(POSITION)', 'CANDIDATE HOLD', 'OFFERED',
                                      "L2 - SELECTED"])
            ).all()
        return jsonify({
            "recruiters": [recruiter.serialize() for recruiter in recruiters],
            "candidates": [candidate.serialize() for candidate in candidates],
            "selected_recruiter": selected_recruiter.serialize() if selected_recruiter else None,
            "selected_candidate": selected_candidate.serialize() if selected_candidate else None,
            "assigned_recruiter": assigned_recruiter.serialize() if assigned_recruiter else None,
            "assignment_message": assignment_message,
            "user_name": user_name
        })

    return jsonify({"error_message": "Unauthorized: You must log in as management user to access this page"}), 401



from flask import jsonify

@app.route('/disable_user', methods=['POST'])
def disable_user():
    data = request.json
    user_id = data['user_id']
    user_status = data['user_status']
    user_name = data['user_name']

    if user_id is None or user_status is None or user_name is None:
        return jsonify({'message': 'User ID, user status, and user name are required'}), 400

    # Find the user making the request
    request_user = User.query.get(user_id)

    if request_user is None or request_user.user_type != 'management':
        return jsonify({'message': 'Unauthorized access'}), 403

    # Find the user to be updated
    user = User.query.filter_by(username=user_name).first()

    if user is None:
        return jsonify({'message': 'User not found'}), 404

    # Change verification status for the user
    user.is_verified = user_status

    # If the user is a recruiter, change verification status for management user with the same username
    if user.user_type == 'recruiter':
        management_user = User.query.filter_by(username=user_name, user_type='management').first()
        if management_user:
            management_user.is_verified = user_status

    try:
        db.session.commit()
        # Return different messages based on user_type
        if user.user_type == 'management':
            if user_status:
                return jsonify({'message': 'Verification status updated for management account'}), 200
            else:
                return jsonify({'message': 'Verification status updated to unverified for management account'}), 200
        elif user.user_type == 'recruiter':
            if user_status:
                return jsonify({'message': 'Verification status updated for recruiter account'}), 200
            else:
                return jsonify({'message': 'Verification status updated to unverified for recruiter account'}), 200
    except Exception as e:
        # Log the exception or return an error message
        db.session.rollback()
        return jsonify({'message': 'Failed to update verification status'}), 500

@app.route('/active_users', methods=['POST'])
def update_user_status():
    data = request.json
    username = data.get('user_name')
    new_status = data.get('new_status')

    try:
        user = User.query.filter_by(username=username).first()
        if user:
            # user.is_verified = new_status
            # db.session.commit()

            # Fetch updated active users list
            active_users_manager = User.query.filter_by(user_type='management').all()
            active_users_manager = sorted(active_users_manager, key=lambda user: user.id)
            active_users_recruiter = User.query.filter_by(user_type='recruiter').all()
            active_users_recruiter = sorted(active_users_recruiter, key=lambda user: user.id)

            return jsonify({
                "message": "User status updated successfully",
                "username": username,
                "active_users_manager": [user.serialize() for user in active_users_manager],
                "active_users_recruiter": [user.serialize() for user in active_users_recruiter]
            })
        else:
            return jsonify({"message": "User not found"}), 404
    except Exception as e:
        db.session.rollback()  # Rollback changes in case of error
        return jsonify({"message": "Error updating user status", "error": str(e)}), 500
        

from flask import jsonify

@app.route('/deactivate_user', methods=['POST'])
def deactivate_user():
    data = request.json
    management_user_id = data.get('user_id')
    recruiter_username = data.get('user_name')
    user_status = data.get('user_status')

    if management_user_id and recruiter_username:  
        # Find the management user
        management_user = User.query.get(management_user_id)

        if management_user and management_user.user_type == 'management':
            # Find the recruiter user by username
            recruiter_user = User.query.filter_by(username=recruiter_username, user_type='recruiter').first()

            if recruiter_user:
                # Change active status for the recruiter user
                recruiter_user.is_verified = user_status
                db.session.commit()

                # Get all user records
                all_users = User.query.all()
                
                # Construct response data
                user_data = [{'id': user.id, 'username': user.username, 'is_active': user.is_verified} for user in all_users]

                if user_status:
                    return jsonify({'message': f'Recruiter account {recruiter_username} has been successfully activated.', 'users': user_data})
                else:
                    return jsonify({'message': f'Recruiter account {recruiter_username} has been successfully deactivated.', 'users': user_data})
            else:
                return jsonify({'message': 'Recruiter user not found or not a recruiter user'})
        else:
            return jsonify({'message': 'Management user not found or not a management user'})
    else:
        return jsonify({'message': 'Both management_user_id and recruiter_username are required'})


        
# @app.route('/verify_checkbox', methods=['POST'])
# def verify_checkbox():
#     data = request.json
#     user_id = data.get('userId')
#     checked = data.get('checked')
#     user = User.query.get(user_id)
#     user.is_verified = checked
#     db.session.commit()
#     return redirect(url_for('active_users'))

import hashlib
from flask_mail import Message

@app.route('/change_password', methods=['POST'])
def change_password():
    data = request.json
    
    if data:
        user_id = data.get('user_id')
        user = User.query.filter_by(id=user_id).first()

        if user:
            user_name = user.username
            user_type = user.user_type

            username = data.get('username')
            old_password = data.get('old_password')
            new_password = data.get('new_password')
            confirm_password = data.get('confirm_password')

            if username == user_name:
                # Check if the provided old password matches the one stored in the database
                hashed_old_password = hashlib.sha256(old_password.encode()).hexdigest()
                if user.password == hashed_old_password:
                    if new_password == confirm_password:
                        # Hash the new password before storing it in the database
                        hashed_new_password = hashlib.sha256(new_password.encode()).hexdigest()
                        user.password = hashed_new_password
                        db.session.commit()

                        # Send the password change notification email
                        msg = Message('Password Changed', sender='ganesh.s@makonissoft.com', recipients=[user.email])
                        msg.body = f'Hello {user.name},\n\nYour password has been successfully changed. Here are your updated credentials:\n\nUsername: {user.username}\nPassword: {new_password}'
                        mail.send(msg)

                        if user_type == 'management':
                            return jsonify({"message": "Password changed successfully for management user."})
                        else:
                            return jsonify({"message": "Password changed successfully for regular user."})
                    else:
                        return jsonify({"message": "New password and confirm password do not match."}) 
                else:
                    return jsonify({"message": "Invalid old password."})
            else:
                return jsonify({"message": "Logged in user does not match the provided username."})
        else:
            return jsonify({"message": "User not found."})
    else:
        return jsonify({"message": "No JSON data provided."})

    # return jsonify({"error": "Unauthorized: You must log in to access this page"})


@app.route('/delete_job_post_message/<int:job_id>')
def delete_job_post_message(job_id):
    job_post = JobPost.query.get(job_id)
    id = job_post.id
    client = job_post.client
    role = job_post.role
    return redirect(url_for('view_all_jobs',client=client,role=role,id=id))

@app.route('/delete_job_post/<int:job_id>', methods=['POST'])
def delete_job_post(job_id):
    # data=request.json
    # job_id=data['job_id']
    job_post = JobPost.query.get(job_id)
    if job_post:
        JobPost.query.filter_by(id=job_id).delete()
        db.session.commit()
        return jsonify({"message": "Job Post Deleted Successfully"}), 200
    else:
        return jsonify({"error": "Job Post not found"}), 404

@app.route('/download_jd/<int:job_id>')
def download_jd(job_id):
    jobpost = JobPost.query.get(job_id)
    if jobpost is None or jobpost.jd_pdf is None:
        return redirect(url_for('dashboard'))

    jd_file = io.BytesIO(jobpost.jd_pdf)
    is_pdf = jd_file.getvalue().startswith(b"%PDF")
    if is_pdf : 
        jd_filename = f"{jobpost.client}_jd.pdf"  # Set the filename as desired
        jd_path = os.path.join(app.config['UPLOAD_FOLDER'], jd_filename)
        with open(jd_path, 'wb') as file:
            file.write(jobpost.jd_pdf)

        # Send the saved resume file for download
        return send_file(jd_path, as_attachment=True)
    else:
        jd_filename = f"{jobpost.client}_jd.docx"  # Set the filename as desired
        jd_path = os.path.join(app.config['UPLOAD_FOLDER'], jd_filename)
        with open(jd_path, 'wb') as file:
            file.write(jobpost.jd_pdf)

        # Send the saved resume file for download
        return send_file(jd_path, as_attachment=True)


import base64
import io
from flask import send_file

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return 'Job post not found', 404  # Return 404 Not Found status
#     # Decode the base64 encoded resume data
#     print("jobpost.jd_pdf",jobpost.jd_pdf.tobytes())
#     if "==" not in str(jobpost.jd_pdf.tobytes()):
#         if request.args.get('decode') == 'base64':
#             # Decode the base64 encoded resume data
#             decoded_resume = base64.b64decode(jobpost.jd_pdf)
#             resume_binary = decoded_resume
#         else:
#             # Retrieve the resume binary data from the database
#             resume_binary = jobpost.jd_pdf.tobytes()  # Convert memoryview to bytes

#         # Determine the mimetype based on the file content
#         is_pdf = resume_binary.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             io.BytesIO(resume_binary),
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         decoded_resume = base64.b64decode(jobpost.jd_pdf)
#         # Create a file-like object (BytesIO) from the decoded resume data
#         resume_file = io.BytesIO(decoded_resume)
#         # Determine the mimetype based on the file content
#         is_pdf = decoded_resume.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             resume_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the job post data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return 'Job post not found', 404  # Return 404 Not Found status

#     # Convert the memoryview to bytes
#     jd_pdf_data = jobpost.jd_pdf.tobytes()
#     if request.args.get('decode') == 'base64':
#         # Decode the base64 encoded data if requested
#         try:
#             jd_pdf_binary = base64.b64decode(jd_pdf_data)
#         except Exception as e:
#             return f'Error decoding base64 data: {e}', 400
#     else:
#         jd_pdf_binary = jd_pdf_data

#     # Determine the mimetype based on the file content
#     is_pdf = jd_pdf_binary.startswith(b"%PDF")
#     mimetype = 'application/pdf' if is_pdf else 'application/msword'

#     # Send the file as a response
#     return send_file(
#         io.BytesIO(jd_pdf_binary),
#         mimetype=mimetype,
#         as_attachment=False
#     )

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the job post data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return jsonify({'error': 'Job post not found'}), 404  # Return JSON error with 404 status

#     # Retrieve the job description data (assumed to be in jobpost.jd_pdf)
#     jd_pdf_data = jobpost.jd_pdf.tobytes()  # Convert memoryview to bytes

#     # Check if the request has 'decode=base64' argument
#     if request.args.get('decode') == 'base64':
#         # Decode the base64 encoded job description data
#         jd_pdf_data = base64.b64decode(jd_pdf_data)

#     # Create a file-like object (BytesIO) from the job description data
#     jd_pdf_file = io.BytesIO(jd_pdf_data)

#     # Determine the mimetype based on the file content
#     mimetype = 'application/pdf' if jd_pdf_data.startswith(b"%PDF") else 'application/msword'

#     # Send the file as a response
#     return send_file(jd_pdf_file, mimetype=mimetype, as_attachment=False)

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the resume data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return 'Job post not found', 404  # Return 404 Not Found status
#     # Decode the base64 encoded resume data
#     print("jobpost.jd_pdf",jobpost.jd_pdf.tobytes())
#     if "==" not in str(jobpost.jd_pdf.tobytes()):
#         if request.args.get('decode') == 'base64':
#             # Decode the base64 encoded resume data
#             decoded_jd_pdf = base64.b64decode(jobpost.jd_pdf)
#             jd_pdf_binary = decoded_jd_pdf
#         else:
#             # Retrieve the resume binary data from the database
#             jd_pdf_binary = jobpost.jd_pdf.tobytes()  # Convert memoryview to bytes
#             print("jd_pdf_binary :",jd_pdf_binary)
#         # Determine the mimetype based on the file content
#         is_pdf = jd_pdf_binary.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             io.BytesIO(jd_pdf_binary),
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         decoded_jd_pdf = base64.b64decode(jobpost.jd_pdf)
#         # Create a file-like object (BytesIO) from the decoded resume data
#         jd_pdf_file = io.BytesIO(decoded_jd_pdf)
#         # Determine the mimetype based on the file content
#         is_pdf = decoded_jd_pdf.startswith(b"%PDF")
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'

#         # Send the file as a response
#         return send_file(
#             jd_pdf_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )

# @app.route('/view_jd/<int:job_id>', methods=['GET'])
# def view_jd(job_id):
#     # Retrieve the job post data from the database using SQLAlchemy
#     jobpost = JobPost.query.filter_by(id=job_id).first()
#     if not jobpost:
#         return 'Job post not found', 404  # Return 404 Not Found status
    
#     # Check if the job post contains a JD PDF
#     if jobpost.jd_pdf:
#         # Decode the base64 string back to its original binary data
#         jd_binary = base64.b64decode(jobpost.jd_pdf)
 
#         # Create a file-like object (BytesIO) from the decoded binary data
#         jd_file = io.BytesIO(jd_binary)
 
#         # Check if the file is a PDF
#         is_pdf = jd_binary.startswith(b"%PDF")
 
#         # Determine the mimetype based on the file type
#         mimetype = 'application/pdf' if is_pdf else 'application/msword'
 
#         # Return the file as a response
#         return send_file(
#             jd_file,
#             mimetype=mimetype,
#             as_attachment=False
#         )
#     else:
#         return 'JD PDF not available', 404  # Return 404 Not Found status if JD PDF is not available

import magic

@app.route('/view_jd/<int:job_id>', methods=['GET'])
def view_jd(job_id):
    # Retrieve the job post data from the database using SQLAlchemy
    jobpost = JobPost.query.filter_by(id=job_id).first()
    if not jobpost:
        return 'Job post not found', 404  # Return 404 Not Found status
    
    # Check if the job post contains a JD PDF
    if jobpost.jd_pdf:
        # Decode the base64 string back to its original binary data
        jd_binary = base64.b64decode(jobpost.jd_pdf)
 
        # Create a file-like object (BytesIO) from the decoded binary data
        jd_file = io.BytesIO(jd_binary)
 
        # Detect the mimetype using more complex logic
        mimetype = detect_mimetype(jd_binary)
 
        # Return the file as a response
        return send_file(
            jd_file,
            mimetype=mimetype,
            as_attachment=False
        )
    else:
        return 'JD PDF not available', 404  # Return 404 Not Found status if JD PDF is not available

def detect_mimetype(data):
    # Check for PDF magic number
    if data.startswith(b"%PDF"):
        return 'application/pdf'
    # Check for MS Word magic number
    elif data.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"):
        return 'application/msword'
    # Add more checks for other file types if needed
    else:
        # If no specific type is detected, fallback to using python-magic
        mime = magic.Magic(mime=True)
        return mime.from_buffer(data)


from flask import Flask, jsonify, request
from datetime import datetime
import pandas as pd

@app.route('/generate_excel', methods=['POST'])
def generate_excel():
    data = request.json

    if not data:
        return jsonify({'error': 'No JSON data provided'})

    user_id = data.get('user_id')
    from_date_str = data.get('from_date')
    to_date_str = data.get('to_date')
    recruiter_names = data.get('recruiter_names', [])

    if not recruiter_names:
        return jsonify({'error': 'Please select any Recruiter'})

    try:
        from_date = datetime.strptime(from_date_str, "%Y-%m-%d")
        to_date = datetime.strptime(to_date_str, "%Y-%m-%d")
    except ValueError:
        return jsonify({'error': 'Invalid date format. Please use YYYY-MM-DD format.'})

    complete_data = [{"recruiter": recruiter_name, "date_created": date.strftime("%Y-%m-%d")} for recruiter_name in recruiter_names for date in pd.date_range(from_date, to_date)]

    complete_df = pd.DataFrame(complete_data)

    # Filter out rows with invalid date strings
    complete_df = complete_df[complete_df['date_created'] != "0"]

    merged_df = pd.concat([pd.DataFrame(data), complete_df]).fillna(0)

    grouped = merged_df.groupby(['recruiter', 'date_created']).size().reset_index(name='count')
    grouped['date_created'] = pd.to_datetime(grouped['date_created'], errors='coerce', format="%Y-%m-%d")
    grouped = grouped.dropna(subset=['date_created'])  # Remove rows with invalid dates
    grouped['date_created'] = grouped['date_created'].dt.strftime("%Y-%m-%d")

    pivot_table = grouped.pivot_table(index='recruiter', columns='date_created', values='count', aggfunc='sum',
                                      fill_value=0, margins=True, margins_name='Grand Total')

    styled_pivot_table = pivot_table.copy()

    recruiters = list(set(recruiter_names))

    styled_pivot_table_json = styled_pivot_table.to_json()

    return jsonify({
        'recruiters': recruiters,
        'styled_pivot_table': styled_pivot_table_json,
        'user_id': user_id,
        'from_date_str': from_date_str,
        'to_date_str': to_date_str
    })

def re_send_notification(recruiter_email, job_id):
    msg = Message('Job Update Notification', sender='ganesh.s@makonissoft.com', recipients=[recruiter_email])
    msg.body = f'Hello,\n\nThe job post with ID {job_id} has been updated.\n\nPlease check your dashboard for more details.'
    mail.send(msg)
   
@app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
def edit_job_post(job_post_id):
    try:
        # Accessing the JSON data from the request
        data = request.json
        user_id = data.get('user_id')
        
        # Retrieve the user
        user = User.query.filter_by(id=user_id).first()
        
        # Check if the user exists and has the right permissions
        if user and user.user_type == 'management':
            # Retrieve the job post to be edited
            job_post = JobPost.query.get(job_post_id)
            
            if job_post:
                # Update job post fields
                job_post.client = data.get('client', job_post.client)
                job_post.experience_min = data.get('experience_min', job_post.experience_min)
                job_post.experience_max = data.get('experience_max', job_post.experience_max)
                job_post.budget_min = data.get('budget_min', job_post.budget_min)
                job_post.budget_max = data.get('budget_max', job_post.budget_max)
                job_post.location = data.get('location', job_post.location)
                job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
                job_post.notice_period = data.get('notice_period', job_post.notice_period)
                job_post.role = data.get('role', job_post.role)
                job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
                job_post.mode = data.get('mode', job_post.mode)
                job_post.job_status = data.get('job_status', job_post.job_status)
                job_post.job_type = data.get('Job_Type', job_post.job_type)  # Updated key 'job_type' to 'Job_Type'
                job_post.skills = data.get('skills', job_post.skills)
                job_post.jd_pdf = data.get('jd_pdf', job_post.jd_pdf)
                job_post.recruiter=data.get('recruiter',job_post.recruiter)

                # Update data_updated_date and data_updated_time
                # current_datetime = datetime.now()
                current_datetime = datetime.now(pytz.timezone('Asia/Kolkata')) 
                job_post.data_updated_date = current_datetime.date()
                job_post.data_updated_time = current_datetime.time()
                
                # Update job post in the database
                db.session.commit()
                
                # Increment num_notification count by 1 for each notification associated with the job post
                notifications = Notification.query.filter_by(job_post_id=job_post_id).all()
                for notification in notifications:
                    notification.num_notification += 1
                
                db.session.commit()
                
                # Return success message
                return jsonify({"message": "Job post updated successfully"}), 200
            else:
                return jsonify({"error": "Job post not found"}), 404
        else:
            return jsonify({"error": "Unauthorized"}), 401
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/jobs_notification/<int:user_id>', methods=['GET'])
def get_jobs_notification(user_id):
    # Retrieve the user
    user = User.query.filter_by(id=user_id).first()
    
    # Check if the user exists and has the right permissions
    if user and user.user_type == 'recruiter':
        recruiter_name = user.username
        
        # Retrieve the notifications for the recruiter where num_notification >= 1
        notifications = Notification.query.filter_by(recruiter_name=recruiter_name).filter(Notification.num_notification >= 1).all()
        
        # Format the notifications as a list of dictionaries
        notifications_list = [
            {
                # 'id': notification.id,
                'job_post_id': notification.job_post_id,
                'recruiter_name': notification.recruiter_name,
                'notification_status': notification.notification_status,
                'num_notification': notification.num_notification
            } for notification in notifications
        ]
        
        return jsonify(notifications_list), 200
    else:
        return jsonify({'error': 'User not found or does not have the right permissions'}), 404

@app.route('/checked_jobs_notification/<int:user_id>', methods=['POST'])
def checked_jobs_notification(user_id):
    data = request.json
    checked_notification_status = data.get('checked_notification_status')

    user = User.query.filter_by(id=user_id).first()
    
    # Check if the user exists and has the right permissions
    if user and user.user_type == 'recruiter':
        recruiter_name = user.username
        
        # Retrieve the notifications for the recruiter
        notifications = Notification.query.filter_by(recruiter_name=recruiter_name).all()
        
        if checked_notification_status:
            # Update the num_notification to 0 for each notification
            for notification in notifications:
                notification.notification_status = checked_notification_status
                notification.num_notification = 0
                db.session.commit()
        
        # Format the notifications as a list of dictionaries
        notifications_list = [
            {
                'id': notification.id,
                'job_post_id': notification.job_post_id,
                'recruiter_name': notification.recruiter_name,
                'notification_status': notification.notification_status,
                'num_notification': notification.num_notification
            } for notification in notifications
        ]
        
        return jsonify(notifications_list), 200
    else:
        return jsonify({'error': 'User not found or does not have the right permissions'}), 404
    
@app.route('/get_candidate_data')
def get_candidate_data():
    candidates = Candidate.query.all()
    candidate_data = []
    for candidate in candidates:
        candidate_data.append({
            'id': candidate.id,
            'name': candidate.name,
            'email': candidate.email,
            'client': candidate.client,
            'current_company':candidate.current_company,
            'position': candidate.position,
            'profile': candidate.profile,
            'current_job_location':candidate.current_job_location,
            'preferred_job_location':candidate.preferred_job_location,
            'skills':candidate.skills,
            'status':candidate.status,
        })
    return jsonify(candidate_data)


@app.route('/send_email', methods=['POST'])
def send_email():
    recipient_email = request.form.get('recipient_email')

    if not recipient_email:
        flash('Recipient email is required.', 'error')
        return redirect(url_for('careers'))

    # Create a link to the page you want to send
    page_link = 'http://127.0.0.1:5001/careers'  # Replace with the actual link

    # Create the email content with a hyperlink
    email_content = f"Click the link below to view active job posts: <a href='{page_link}'>{page_link}</a>"

    # Create an email message
    message = Message('Active Job Posts', sender='ganesh.s@makonissoft.com', recipients=[recipient_email])
    message.html = email_content

    # Send the email
    mail.send(message)

    flash('Email sent successfully!', 'success')
    return redirect(url_for('careers'))

#new
@app.route('/careers', methods=['GET'])
def careers():
    user_type = session.get('user_type', None)
    is_logged_in = 'user_id' in session
    candidate_message = request.args.get('candidate_message')
    print(candidate_message)

    # Query the database to retrieve active job posts and sort them by date_created in descending order
    active_jobs = JobPost.query.filter_by(job_status='Active').order_by(JobPost.date_created.desc()).all()

    return render_template('careers.html', jobs=active_jobs, user_type=user_type, is_logged_in=is_logged_in,candidate_message=candidate_message)

#new
@app.route('/apply_careers', methods=['GET', 'POST'])
def apply_careers():
    user_id = session.get('user_id')
    if not user_id:
        # User is not authenticated, you can redirect them to a login page or take appropriate action
        return redirect(url_for('career_login'))
    user = Career_user.query.get(user_id)
    if request.method == 'GET':
        job_id = request.args.get('job_id')
        client = request.args.get('client')
        profile = request.args.get('profile')
        name = user.name
        email = user.email

        if job_id:
            matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Hold')).first()
            if matching_job_post:
                return render_template('job_on_hold.html')
        
        job_post = JobPost.query.get(job_id)
        experience_min = job_post.experience_min

        job_ids = db.session.query(JobPost.id).filter(JobPost.client == client, JobPost.job_status == 'Active').all()
        job_roles = db.session.query(JobPost.role).filter(JobPost.client == client).all()

        ids = [job_id[0] for job_id in job_ids]
        roles = [job_role[0] for job_role in job_roles]

        candidate_data = None
        if 'candidate_data' in request.args:
            candidate_data = ast.literal_eval(request.args['candidate_data'])

        return render_template('apply_careers.html', candidate_data=candidate_data, job_id=job_id,
                               client=client, profile=profile, ids=ids, roles=roles,
                               name=name, email=email,experience_min=experience_min)

    if request.method == 'POST':
        try:
            job_id = request.form['job_id']
            name = request.form['name']
            mobile = request.form['mobile']
            email = request.form['email']
            client = request.form['client']
            profile = request.form['profile']
            skills = request.form['skills']

            # Ensure client and job_id are integers
            job_id = int(job_id)

            # Check if the job post is active
            matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Active')).first()
            if not matching_job_post:
                return render_template('job_on_hold.html')

            # Handle other form fields...
            current_company = request.form['current_company']
            position = request.form['position']
            current_job_location = request.form['current_job_location']
            preferred_job_location = request.form['preferred_job_location']
            qualifications = request.form['qualifications']
            experience = request.form['experience']
            exp_months = request.form['exp_months']
            experience = experience + '.' + exp_months
            relevant_experience = request.form['relevant_experience']
            relevant_exp_months = request.form['relevant_exp_months']
            relevant_experience = relevant_experience + '.' + relevant_exp_months
            currency_type_current = request.form['currency_type_current']
            currency_type_except = request.form['currency_type_except']
            current_ctc = currency_type_current + " " + request.form['current_ctc']
            expected_ctc = currency_type_except + " " + request.form['expected_ctc']
            linkedin = request.form['linkedin']

            # Handle file upload
            filename = None
            resume_binary = None
            if 'resume' in request.files:
                resume_file = request.files['resume']
                if resume_file and allowed_file(resume_file.filename):
                    # Convert the resume file to binary data
                    resume_binary = resume_file.read()
                    filename = secure_filename(resume_file.filename)
                else:
                    return render_template('apply_careers.html', error_message='Invalid file extension')

            notice_period = request.form['notice_period']
            last_working_date = None
            buyout = False
            period_of_notice = None

            if notice_period == 'yes':
                last_working_date = request.form['last_working_date']
                buyout = 'buyout' in request.form
            elif notice_period == 'no':
                period_of_notice = request.form['months']
                buyout = 'buyout' in request.form
            elif notice_period == 'completed':
                last_working_date = request.form['last_working_date']

            holding_offer = request.form['holding_offer']

            if holding_offer == 'yes':
                total = request.form['total']
                if total == '':
                    total = 0
                else:
                    total = int(request.form['total'])
                package_in_lpa = request.form['package_in_lpa']
                if package_in_lpa == '':
                    package_in_lpa = 0
                else:
                    package_in_lpa = float(request.form['package_in_lpa'])
            else:
                total = None
                package_in_lpa = None

            reason_for_job_change = request.form.get('reason_for_job_change')
            remarks = request.form.get('remarks')

            reference = request.form['reference']
            reference_name = None
            reference_position = None
            reference_information = None

            if reference == 'yes':
                reference_name = request.form['reference_name']
                reference_position = request.form['reference_position']
                reference_information = request.form['reference_information']
            elif reference == 'no':
                reference_name = None
                reference_position = None
                reference_information = None

            existing_candidate = Candidate.query.filter(
                and_(Candidate.profile == profile, Candidate.client == client, Candidate.email == email,
                     Candidate.mobile == mobile)).first()
            if existing_candidate:
                return render_template('candidate_exists.html',
                                       error_message='Candidate with the same profile and client already exists')

            # Create a new Candidate object
            new_candidate = Candidate(
                job_id=job_id,
                name=name,
                mobile=mobile,
                email=email,
                client=client,
                current_company=current_company,
                position=position,
                profile=profile,
                resume=resume_binary,
                current_job_location=current_job_location,
                preferred_job_location=preferred_job_location,
                qualifications=qualifications,
                experience=experience,
                relevant_experience=relevant_experience,
                current_ctc=current_ctc,
                expected_ctc=expected_ctc,
                notice_period=notice_period,
                last_working_date=last_working_date if notice_period == 'yes' or notice_period == 'completed' else None,
                buyout=buyout,
                holding_offer=holding_offer,
                total=total,
                package_in_lpa=package_in_lpa,
                linkedin_url=linkedin,
                reason_for_job_change=reason_for_job_change,
                status='None',
                remarks=remarks,
                skills=skills,
                period_of_notice=period_of_notice,
                reference=reference,
                reference_name=reference_name,
                reference_position=reference_position,
                reference_information=reference_information
            )

            new_candidate.date_created = date.today()
            new_candidate.time_created = datetime.now().time()

            # Commit the new candidate to the database
            db.session.add(new_candidate)
            db.session.commit()

            try:
                msg = Message('Successful Submission of Your Job Application', sender='ganesh.s@makonissoft.com', recipients=[email])
                msg.body = f"Dear { name },\n Congratulations! Your job application has been successfully submitted for the position at {client} for the role of {profile}. We appreciate your interest in joining our team.\n\n  Our dedicated recruiter will review your application, and you can expect to hear from us within the next 24 hours.\n\nBest wishes for your application process!\n\n Regards, \n\nTeam\nMakonis Talent Track Pro\nrecruiterpro@makonissoft.com\n"
                mail.send(msg)
            except Exception as e:
                # Handle email sending errors, log the error
                return render_template('error.html', error_message=f"Error sending thank-you email: {str(e)}")

            return redirect(url_for('careers', candidate_message='Candidate Added Successfully'))

        except Exception as e:
            # Handle any exceptions here (e.g., log the error, return an error page)
            return render_template('error.html', error_message=str(e))

    return redirect(url_for('careers'))


#new
# User Login
@app.route('/career_login', methods=['GET', 'POST'])
def career_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = Career_user.query.filter_by(username=username, password=password).first()

        if user:
            # Store the user's session or token
            session['user_id'] = user.id
            return redirect(url_for('careers'))

    return render_template('career_login.html')

#new
@app.route('/career_logout')
def career_logout():
    # Clear the user's session
    session.pop('user_id', None)
    return redirect(url_for('careers'))

#new
@app.route('/career_register', methods=['GET', 'POST'])
def career_register():
    if request.method == 'POST':
        username = request.form['username']
        name = request.form['name']
        email = request.form['email']
        password = request.form['password']

        # Create a new user and add it to the database
        new_user = Career_user(username=username, name=name, email=email, password=password)
        db.session.add(new_user)
        db.session.commit()

        return redirect(url_for('career_login'))

    return render_template('career_registration.html')

#new
@app.route('/career_dashboard')
def career_dashboard():
    edit_candidate_message = request.args.get('edit_candidate_message')
    if 'user_id' in session and 'user_type' in session:
        page_no = request.args.get('page_no')
        candidate_message = request.args.get('candidate_message')
        signup_message = request.args.get('signup_message')
        job_message = request.args.get('job_message')
        update_candidate_message = request.args.get('update_candidate_message')
        delete_message = request.args.get("delete_message")

        user_id = session['user_id']
        user_type = session['user_type']
        user_name = session['user_name']

        if user_type == 'management':
            users = User.query.all()
            candidates = Candidate.query.filter((Candidate.reference.is_not(None))).all()
            candidates = sorted(candidates, key=lambda candidate: candidate.id)
            JobsPosted = JobPost.query.all()
            # data = json.dumps(candidates, sort_keys=False)
            return render_template('career_dashboard.html', users=users, user_type=user_type, user_name=user_name,
                                   candidates=candidates, update_candidate_message=update_candidate_message,
                                   candidate_message=candidate_message, delete_message=delete_message,
                                   JobsPosted=JobsPosted, signup_message=signup_message, job_message=job_message,
                                   page_no=page_no, edit_candidate_message=edit_candidate_message)
        elif user_type == 'recruiter':
            recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
            recruiter_name = User.query.get(user_id).name
            if recruiter:
                candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name,
                                                         Candidate.reference.is_not(None))).all()
                candidates = sorted(candidates, key=lambda candidate: candidate.id)
                career_count_notification_no = Career_notification.query.filter(Career_notification.notification_status == 'false',
                                                                  Career_notification.recruiter_name == user_name).count()
                career_notifications = Career_notification.query.filter(
                    Career_notification.recruiter_name.contains(recruiter_name)).all()

                for career_notification in career_notifications:
                    if career_notification.notification_status == False:
                        career_notification.notification_status = True
                        db.session.commit()
                return render_template('career_dashboard.html', user=recruiter, user_type=user_type, user_name=user_name,
                                       candidates=candidates, candidate_message=candidate_message,
                                       update_candidate_message=update_candidate_message,
                                       career_count_notification_no=career_count_notification_no,
                                       edit_candidate_message=edit_candidate_message, page_no=page_no)
        else:
            user = User.query.filter_by(id=user_id).first()
            if user:
                candidates = Candidate.query.filter_by(recruiter=user.name).all()  # Filter candidates by user's name
                return render_template('career_dashboard.html', user=user, user_type=user_type, candidates=candidates)

    return redirect(url_for('index'))

#new
@app.route('/website_candidate_assign', methods=['GET', 'POST'])
def website_candidate_assign():
    assignment_message = request.args.get('assignment_message')
    if 'user_id' in session and 'user_type' in session and session['user_type'] == 'management':
        user_name = session['user_name']
        recruiters = User.query.filter_by(user_type='recruiter').all()

        if request.method == 'POST':
            assign_recruiter_id = request.form.get('assign_recruiter_id')
            selected_candidate_ids = request.form.getlist('selected_candidate_ids')

            if assign_recruiter_id and selected_candidate_ids:
                assigned_recruiter = User.query.get(assign_recruiter_id)
                if assigned_recruiter:
                    # Fetch selected candidates by their IDs
                    candidates = Candidate.query.filter(
                        Candidate.id.in_(selected_candidate_ids),
                        Candidate.recruiter.is_(None),
                        Candidate.management.is_(None)
                    ).all()

                    for candidate in candidates:
                        # Assign the selected recruiter to the candidate
                        candidate.recruiter = assigned_recruiter.name
                        # Send an email to the assigned recruiter
                        send_career_email(assigned_recruiter.email, 'Alert! New Candidate Assignment ',
                                          f'Dear {assigned_recruiter.name}\n\n,A new candidate application has been assigned to you. Please access your dashboard to view the details.\n\nCandidate Name: {candidate.name}\n\nClient: {candidate.client}\n\nRole: {candidate.profile}\n\nAssigned by Manager: {user_name}\n\nFeel free to reach out if you have any questions during the recruitment process.\n\nRegards,\n\nTeam\nMakonis Talent Track Pro\nrecruiterpro@makonissoft.com')

                    db.session.commit()

                    # Create notifications for the assigned recruiter
                    notifications = []
                    for candidate in candidates:
                        notification = Career_notification(
                            recruiter_name=assigned_recruiter.name,
                            notification_status=False  # You may set this to True for unread notifications
                        )
                        notifications.append(notification)

                    db.session.add_all(notifications)
                    db.session.commit()

                    return redirect(
                        url_for('website_candidate_assign', assignment_message='Candidates Assigned Successfully'))

        candidates = Candidate.query.filter(
            Candidate.recruiter.is_(None),
            Candidate.management.is_(None)
        ).all()

        candidate_count = Candidate.query.filter(
            Candidate.recruiter.is_(None),
            Candidate.management.is_(None)
        ).count()

        return render_template(
            'website_candidate_assign.html',
            recruiters=recruiters,
            candidates=candidates,
            assignment_message=assignment_message,
            user_name=user_name,
            candidate_count=candidate_count
        )

    return redirect(url_for('index'))

#new
def send_career_email(to, subject, message):
    msg = Message(subject, sender='ganesh.s@makonissoft.com', recipients=[to])
    msg.body = message
    mail.send(msg)

####################################################################################################################################

import base64
import io
import re
from flask import Flask, request, jsonify
import fitz  # PyMuPDF
from docx import Document

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(file):
    """
    Extract text from PDF or DOCX files.
    
    Parameters:
        file (BytesIO): File-like object.
    
    Returns:
        str: Extracted text.
    """
    try:
        file.seek(0)
        header = file.read(4)
        file.seek(0)
        if header.startswith(b'%PDF'):
            return extract_text_from_pdf(file)
        elif header.startswith(b'PK\x03\x04'):
            return extract_text_from_docx(file)
        else:
            return ""  # Unsupported file format
    except Exception as e:
        print(f"Error determining file type: {e}")
        return ""

def extract_text_from_pdf(file):
    """
    Extract text from a PDF file.
    
    Parameters:
        file (BytesIO): PDF file-like object.
    
    Returns:
        str: Extracted text.
    """
    text = ""
    try:
        with fitz.open(stream=file, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(file):
    """
    Extract text from a DOCX file.
    
    Parameters:
        file (BytesIO): DOCX file-like object.
    
    Returns:
        str: Extracted text.
    """
    text = ""
    try:
        doc = Document(file)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
    return text

def extract_skills_from_resume(text, skills_list):
    found_skills = [skill for skill in skills_list if skill.lower() in text.lower()]
    return found_skills

def extract_email(text):
    email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    email_matches = re.findall(email_regex, text)
    return email_matches[-1].rstrip('.,') if email_matches else "No email found"
    
def extract_phone_number(text):
    phone_regex = r'\b\d{10}\b'
    phone_matches = re.findall(phone_regex, text)
    return phone_matches[-1] if phone_matches else "No phone number found"

# def extract_phone_number(text):
#     phone_regex = r'\+?\d[\d -]{8,12}\d'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[-1] if phone_matches else "No phone number found"

def extract_name(text):
    """
    Extract the name from the first few lines of the resume text.
    
    Parameters:
        text (str): Resume text.
    
    Returns:
        str: Extracted name.
    """
    lines = text.split('\n')
    name_words = []  # List to store the words of the name
    
    # Regular expressions to identify lines that are likely contact details
    phone_pattern = re.compile(r'\b(\+?\d[\d\-\.\s]+)?\d{10}\b')
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    
    for line in lines[:5]:  # Look at the first five lines where the name is likely to appear
        # Skip lines that are likely to be contact details
        if phone_pattern.search(line) or email_pattern.search(line):
            continue
        
        # Remove common salutations and titles
        cleaned_line = re.sub(r'\b(Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam)\b', '', line, flags=re.IGNORECASE).strip()
        
        # Extract names with up to three words
        words = cleaned_line.split()
        name_words.extend(words)  # Add words from the current line to the list
        
        if len(name_words) <= 2:
            continue  # Continue accumulating words if we have less than or equal to three words
        else:
            # Stop accumulating if we exceed three words and return the concatenated name
            return ' '.join(word.capitalize() for word in name_words[:3]).rstrip('.,')
    
    # Return the concatenated name if found within the first five lines
    if name_words:
        return ' '.join(word.capitalize() for word in name_words[:3]).rstrip('.,')
    
    return "No name found"

# def extract_name(text):
#     """
#     Extract the name from the first few lines of the resume text.
    
#     Parameters:
#         text (str): Resume text.
    
#     Returns:
#         str: Extracted name.
#     """
#     lines = text.split('\n')
#     name_words = []  # List to store the words of the name
#     for line in lines[:5]:  # Look at the first five lines where the name is likely to appear
#         # Remove common salutations and titles
#         cleaned_line = re.sub(r'Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam', '', line, flags=re.IGNORECASE).strip()
#         # Extract names with up to three words
#         words = cleaned_line.split()
#         name_words.extend(words)  # Add words from the current line to the list
#         if len(name_words) <= 2:
#             continue  # Continue accumulating words if we have less than or equal to three words
#         else:
#             # Stop accumulating if we exceed three words and return the concatenated name
#             return ' '.join(word.capitalize() for word in name_words).rstrip('.,')
#     # Return the concatenated name if found within the first five lines
#     if name_words:
#         return ' '.join(word.capitalize() for word in name_words).rstrip('.,')
#     return "No name found"

# def extract_name(text):
#     """
#     Extract the name from the first few lines of the resume text.
    
#     Parameters:
#         text (str): Resume text.
    
#     Returns:
#         str: Extracted name.
#     """
#     lines = text.split('\n')
#     for line in lines[:5]:  # Look at the first five lines where the name is likely to appear
#         # Remove common salutations and titles
#         cleaned_line = re.sub(r'Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam', '', line, flags=re.IGNORECASE).strip()
#         # Extract names with at least two words
#         words = cleaned_line.split()
#         if len(words) >= 1:
#             # Capitalize the first letter of each word in the name
#             return ' '.join(word.capitalize() for word in words).rstrip('.,')
#     return "No name found"

# def extract_name(text):
#     """
#     Extract the name from the first few lines of the resume text.
    
#     Parameters:
#         text (str): Resume text.
    
#     Returns:
#         str: Extracted name.
#     """
#     lines = text.split('\n')
#     last_name = ""
#     for line in lines:
#         # Remove common salutations and titles
#         cleaned_line = re.sub(r'Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam', '', line, flags=re.IGNORECASE)
#         # Extract names with at least two words
#         words = cleaned_line.split()
#         if len(words) >= 2:
#             # Capitalize the first letter of each word in the name
#             last_name = ' '.join(word.capitalize() for word in words).rstrip('.,')
#     return last_name if last_name else "No name found"

@app.route('/parse_resume', methods=['POST'])
def parse_resume():
    if 'resume' not in request.json:
        return jsonify({"error": "No resume data provided"}), 400
    
    data = request.json
    resume_data = data['resume']
    
    try:
        decoded_resume = base64.b64decode(resume_data)
    except Exception as e:
        return jsonify({"error": "Invalid resume data"}), 400
    
    resume_file = io.BytesIO(decoded_resume)
    resume_text = extract_text(resume_file)
    
    if not resume_text:
        return jsonify({"error": "No text found in the resume data"}), 400

    it_skills = [ 
        'Data Analysis', 'Machine Learning', 'Communication', 'Project Management',
        'Deep Learning', 'SQL', 'Tableau', 'C++', 'C', 'Front End Development', 'JAVA', 
        'Java Full Stack', 'React JS', 'Node JS','Programming (Python, Java, C++)',
        'Data Analysis and Visualization','Artificial Intelligence','Programming',
        'Database Management (SQL)','Web Development (HTML, CSS, JavaScript)',
        'Machine Learning and Artificial Intelligence','Network Administration',
        'Software Development and Testing','Embedded Systems','CAD and 3D Modeling',
        'HTML5', 'CSS3', 'Jquery', 'Bootstrap', 'XML', 'JSON', 'ABAP', 'SAPUI5',
        'Agile Methodology', 'Frontend Development', 'Jira', 'Odata', 'BTP', 'Fiori Launchpad', 
        'Python', 'JavaScript', 'HTML', 'CSS','React', 'Node.js', 'Django', 'Git', 'AWS',
        'Linux','DevOps','Linear Regression','Logistic Regression','Decision Tree',
        'SVM (Support Vector Machine)','Ensembles','Random Forest','Clustering',
        'PCA (Principal Component Analysis)','K-means','Recommendation System',
        'Market Basket Analysis','CNN','RNN','LSTM','Natural Language Processing',
        'NLTK','LGBM','XGBoost','Transformers','Siamese network','BTYD (Buy Till You Die)',
        'ML Ops Tools: Azure Synapse','Azure ML','Azure Databricks','ML flow','Airflow',
        'Kubernetes','Dockers','Data Streaming – Kafka','Flask','LT Spice','Wireshark',
        'Ansys Lumerical','Zemax OpticStudio','Xilinx Vivado','Google Collab','MATLAB'
    ]
    
    non_it_skills = [
        'Communication Skills', 'Teamwork', 'Problem Solving', 'Time Management', 'Leadership',
        'Creativity', 'Adaptability', 'Critical Thinking', 'Analytical Skills', 'Attention to Detail',
        'Customer Service', 'Interpersonal Skills', 'Negotiation Skills', 'Project Management', 
        'Presentation Skills', 'Research Skills', 'Organizational Skills', 'Multitasking',
        'Decision Making', 'Emotional Intelligence', 'Conflict Resolution', 'Networking', 
        'Strategic Planning', 'Public Speaking', 'Writing Skills', 'Sales Skills', 'Marketing', 
        'Finance', 'Human Resources', 'Training and Development', 'Event Planning', 'Language Proficiency',
        'Problem-Solving', 'Sales', 'Marketing', 'Financial Analysis', 'Customer Relationship Management (CRM)', 
        'Quality Management', 'Supply Chain Management', 'Logistics', 'Health and Safety', 'Public Relations', 
        'Social Media Management', 'Content Creation', 'Graphic Design', 'Video Editing', 'Photography', 
        'Data Entry', 'Administrative Support', 'Customer Support'
    ]

    extracted_it_skills = extract_skills_from_resume(resume_text, it_skills)
    extracted_nonit_skills = extract_skills_from_resume(resume_text, non_it_skills)
    non_it_skills_final = list(set(extracted_nonit_skills) - set(extracted_it_skills))

    skills_it = ", ".join(extracted_it_skills) if extracted_it_skills else "No skills found"
    skills_non_it = ", ".join(non_it_skills_final) if non_it_skills_final else "No skills found"

    email_text = extract_email(resume_text)
    phone_text = extract_phone_number(resume_text)
    name_text = extract_name(resume_text)

    return jsonify({
        "name": name_text,
        "mail": email_text,
        "phone": phone_text,
        "skill1": skills_it,
        "skill2": skills_non_it
    })


# def extract_text(file):
#     """
#     Extract text from PDF or DOCX files.
    
#     Parameters:
#         file (BytesIO): File-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     try:
#         file.seek(0)
#         header = file.read(4)
#         file.seek(0)
#         if header.startswith(b'%PDF'):
#             return extract_text_from_pdf(file)
#         elif header.startswith(b'PK\x03\x04'):
#             return extract_text_from_docx(file)
#         else:
#             return ""  # Unsupported file format
#     except Exception as e:
#         print(f"Error determining file type: {e}")
#         return ""

# def extract_text_from_pdf(file):
#     """
#     Extract text from a PDF file.
    
#     Parameters:
#         file (BytesIO): PDF file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         with fitz.open(stream=file, filetype="pdf") as doc:
#             for page in doc:
#                 text += page.get_text()
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#     return text

# def extract_text_from_docx(file):
#     """
#     Extract text from a DOCX file.
    
#     Parameters:
#         file (BytesIO): DOCX file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         doc = Document(file)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#     return text

# def extract_skills_from_resume(text, skills_list):
#     found_skills = [skill for skill in skills_list if skill.lower() in text.lower()]
#     return found_skills

# def extract_email(text):
#     email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#     email_matches = re.findall(email_regex, text)
#     return email_matches[0].rstrip('.,') if email_matches else "No email found"

# def extract_phone_number(text):
#     phone_regex = r'\+?\d[\d -]{8,12}\d'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[0] if phone_matches else "No phone number found"

# def extract_name(text):
#   """
#   Extract the name from the first few lines of the resume text, handling common salutations,
#   titles, single-word names, and punctuation at the end.

#   Args:
#       text (str): Resume text.

#   Returns:
#       str: Extracted name (or "No name found" if no suitable name is found).
#   """

#   lines = text.split('\n')
#   for line in lines[:3]:  # Check only the first 3 lines
#     cleaned_line = re.sub(r'Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam', '', line, flags=re.IGNORECASE)
#     words = cleaned_line.strip().split()

#     # Handle single-word names (e.g., initials)
#     if len(words) == 1:
#       # If it's all uppercase, consider it a name
#       if words[0].isupper():
#         return words[0]
#       else:
#         continue  # Skip single-word names that are not all uppercase

#     # Extract and capitalize name with two or more words
#     if len(words) >= 2:
#       # Remove trailing punctuation (., !) from the last word
#       last_word = words[-1].rstrip('.,!')
#       name = ' '.join(word.capitalize() for word in words[:-1]) + ' ' + last_word
#       return name

#   return "No name found"
    
# @app.route('/parse_resume', methods=['POST'])
# def parse_resume():
#     if 'resume' not in request.json:
#         return jsonify({"error": "No resume data provided"}), 400
    
#     data = request.json
#     resume_data = data['resume']
    
#     try:
#         decoded_resume = base64.b64decode(resume_data)
#     except Exception as e:
#         return jsonify({"error": "Invalid resume data"}), 400
    
#     resume_file = io.BytesIO(decoded_resume)
#     resume_text = extract_text(resume_file)
    
#     if not resume_text:
#         return jsonify({"error": "No text found in the resume data"}), 400

#     it_skills = [ 
#         'Data Analysis', 'Machine Learning', 'Communication', 'Project Management',
#         'Deep Learning', 'SQL', 'Tableau', 'C++', 'C', 'Front End Development', 'JAVA', 
#         'Java Full Stack', 'React JS', 'Node JS','Programming (Python, Java, C++)',
#         'Data Analysis and Visualization','Artificial Intelligence','Programming',
#         'Database Management (SQL)','Web Development (HTML, CSS, JavaScript)',
#         'Machine Learning and Artificial Intelligence','Network Administration',
#         'Software Development and Testing','Embedded Systems','CAD and 3D Modeling',
#         'HTML5', 'CSS3', 'Jquery', 'Bootstrap', 'XML', 'JSON', 'ABAP', 'SAPUI5',
#         'Agile Methodology', 'Frontend Development', 'Jira', 'Odata', 'BTP', 'Fiori Launchpad', 
#         'Python', 'JavaScript', 'HTML', 'CSS','React', 'Node.js', 'Django', 'Git', 'AWS',
#         'Linux','DevOps','Linear Regression','Logistic Regression','Decision Tree',
#         'SVM (Support Vector Machine)','Ensembles','Random Forest','Clustering',
#         'PCA (Principal Component Analysis)','K-means','Recommendation System',
#         'Market Basket Analysis','CNN','RNN','LSTM','Natural Language Processing',
#         'NLTK','LGBM','XGBoost','Transformers','Siamese network','BTYD (Buy Till You Die)',
#         'ML Ops Tools: Azure Synapse','Azure ML','Azure Databricks','ML flow','Airflow',
#         'Kubernetes','Dockers','Data Streaming – Kafka','Flask','LT Spice','Wireshark',
#         'Ansys Lumerical','Zemax OpticStudio','Xilinx Vivado','Google Collab','MATLAB'
#     ]
    
#     non_it_skills = [
#         'Communication Skills', 'Teamwork', 'Problem Solving', 'Time Management', 'Leadership',
#         'Creativity', 'Adaptability', 'Critical Thinking', 'Analytical Skills', 'Attention to Detail',
#         'Customer Service', 'Interpersonal Skills', 'Negotiation Skills', 'Project Management', 
#         'Presentation Skills', 'Research Skills', 'Organizational Skills', 'Multitasking',
#         'Decision Making', 'Emotional Intelligence', 'Conflict Resolution', 'Networking', 
#         'Strategic Planning', 'Public Speaking', 'Writing Skills', 'Sales Skills', 'Marketing', 
#         'Finance', 'Human Resources', 'Training and Development', 'Event Planning', 'Language Proficiency',
#         'Problem-Solving', 'Sales', 'Marketing', 'Financial Analysis', 'Customer Relationship Management (CRM)', 
#         'Quality Management', 'Supply Chain Management', 'Logistics', 'Health and Safety', 'Public Relations', 
#         'Social Media Management', 'Content Creation', 'Graphic Design', 'Video Editing', 'Photography', 
#         'Data Entry', 'Administrative Support', 'Customer Support', 'Teaching', 'Mentoring', 'Coaching', 
#         'Retail Management', 'Hospitality Management', 'Event Management', 'Creative Writing', 'Content Marketing', 
#         'Copywriting', 'Publications', 'Translation', 'Counseling', 'Fitness Instruction', 'Nutrition', 'Wellness', 
#         'Fashion Design', 'Interior Design', 'Artistic Skills', 'Music', 'Sports', 'Culinary Arts', 'Photography', 
#         'Videography', 'Project Coordination', 'Community Outreach', 'Volunteer Management', 'Fundraising', 
#         'Political Campaigning', 'Government Relations', 'Policy Analysis', 'Nonprofit Management', 'Grant Writing', 
#         'Fundraising', 'Event Planning', 'Real Estate', 'Property Management', 'Construction Management', 
#         'Facilities Management', 'Environmental Sustainability', 'Energy Management', 'Public Health', 
#         'Healthcare Administration', 'Nursing', 'Dental Hygiene', 'Pharmacy', 'Physical Therapy', 'Occupational Therapy', 
#         'Social Work', 'Child Care', 'Elderly Care', 'Counseling', 'Psychology', 'Sociology', 'Anthropology', 'Archaeology', 
#         'Geography', 'History', 'Political Science', 'Economics', 'Philosophy', 'Theology', 'Linguistics', 'Literature', 
#         'Creative Writing', 'Journalism', 'Broadcasting', 'Public Relations', 'Marketing', 'Advertising', 'Market Research', 
#         'Retail Sales', 'Wholesale Sales', 'Account Management', 'Client Relations', 'Customer Service', 'Conflict Resolution', 
#         'Presentation Skills', 'Public Speaking', 'Writing', 'Editing', 'Proofreading', 'Content Creation', 'Graphic Design', 
#         'Visual Merchandising', 'Retail Operations', 'Inventory Management', 'Supply Chain', 'Logistics', 'Quality Assurance', 
#         'Process Improvement', 'Project Management', 'Financial Planning', 'Budgeting', 'Financial Analysis', 'Bookkeeping', 
#         'Data Entry', 'Administrative Support', 'Executive Assistance', 'Time Management', 'Organizational Skills', 'Event Planning', 
#         'Event Coordination', 'Event Marketing', 'Catering', 'Venue Management', 'Wedding Planning', 'Trade Show Coordination', 
#         'Customer Service', 'Conflict Resolution', 'Problem-Solving', 'Decision Making', 'Team Collaboration', 'Leadership', 
#         'Supervision', 'Employee Training', 'Performance Management', 'Recruitment', 'Human Resources', 'Payroll Administration', 
#         'Employee Relations', 'Safety Compliance', 'Labor Relations', 'Legal Compliance', 'Contract Negotiation', 'Risk Management', 
#         'Policy Development', 'Quality Management', 'Process Improvement', 'Supply Chain Management', 'Logistics', 'Inventory Control', 
#         'Procurement', 'Distribution', 'Quality Assurance', 'Process Improvement', 'Product Development', 'Marketing Strategy', 
#         'Brand Management', 'Market Research', 'Public Relations', 'Social Media Management', 'Content Creation', 'Copywriting', 
#         'Email Marketing', 'Sales Strategy', 'Client Relationship Management', 'Sales Forecasting', 'Lead Generation', 
#         'Account Management', 'Customer Retention', 'Sales Presentations', 'Networking', 'Public Speaking', 'Team Collaboration', 
#         'Project Management', 'Client Communications', 'Technical Support', 'Troubleshooting', 'Network Administration', 
#         'Quality Assurance', 'Project Management', 'Technical Writing', 'Documentation', 'Research and Development', 'Innovation', 
#         'Problem-Solving', 'Critical Thinking', 'Attention to Detail', 'Collaboration', 'Time Management', 'Adaptability', 'Leadership', 
#         'Creativity', 'Analytical Skills', 'Data Analysis', 'Statistical Analysis', 'Mathematics', 'Physics', 'Chemistry', 'Biology', 
#         'Geology', 'Environmental Science', 'Meteorology', 'Agricultural Science', 'Animal Science', 'Food Science', 'Nutrition', 
#         'Dietetics', 'Physical Therapy', 'Occupational Therapy', 'Speech-Language Pathology', 'Nursing', 'Pharmacy', 'Dentistry', 
#         'Veterinary Medicine', 'Medical Research', 'Medical Writing', 'Clinical Trials', 'Epidemiology', 'Public Health', 
#         'Healthcare Administration', 'Health Informatics', 'Fitness Instruction', 'Nutrition Counseling', 'Wellness Coaching', 
#         'Yoga Instruction', 'Personal Training', 'Physical Education', 'Sports Coaching', 'Athletic Training', 'Recreation', 
#         'Dance Instruction', 'Music Instruction', 'Art Instruction', 'Photography', 'Video Editing', 'Graphic Design', 'Interior Design', 
#         'Fashion Design', 'Culinary Arts', 'Baking', 'Cooking', 'Restaurant Management', 'Hotel Management', 'Tourism', 'Event Planning', 
#         'Museum Management', 'Library Science', 'Archiving', 'Curatorial Work', 'Conservation', 'Environmental Science', 
#         'Sustainability', 'Renewable Energy', 'Climate Change', 'Environmental Policy', 'Wildlife Conservation', 'Forestry', 
#         'Natural Resource Management', 'Ecology', 'Geography', 'Urban Planning', 'Civil Engineering', 'Structural Engineering', 
#         'Transportation Engineering', 'Geotechnical Engineering', 'Environmental Engineering', 'Water Resources Engineering', 
#         'Surveying', 'Architecture', 'Landscape Architecture', 'Interior Design', 'Urban Design', 'Real Estate Development', 
#         'Property Management', 'Construction Management', 'Building Inspection', 'Facilities Management', 'Space Planning', 
#         'Urban Planning', 'Public Administration', 'Policy Analysis', 'Government Relations', 'Political Campaigning', 'Public Policy', 
#         'Economics', 'Finance', 'Accounting', 'Actuarial Science', 'MS Office','Powerpoint','ms word','ms excel' 
#     ]

#     extracted_skills = extract_skills_from_resume(resume_text, it_skills)
#     extracted_nonit_skills = extract_skills_from_resume(resume_text, non_it_skills)
#     non_it_skills = list(set(extracted_nonit_skills) - set(extracted_skills))

#     skills_it = ", ".join(extracted_skills) if extracted_skills else "No skills found"
#     skills_non_it = ", ".join(non_it_skills) if non_it_skills else "No skills found"

#     mail_text = extract_email(resume_text)
#     phone_text = extract_phone_number(resume_text)

#     first_line = resume_text.split('\n')[0]
#     words = first_line.split()
#     first_3_words = ' '.join(words[:3])

#     name_text = "No name found"
#     if "RESUME" in first_line or "Resume" in first_line or "BIODATA" in first_line or "BioData" in first_line or "biodata" in first_line:
#         second_line = resume_text.split('\n')[1] if len(resume_text.split('\n')) > 1 else ""
#         words1 = second_line.split()
#         first_5_words_in_2 = ' '.join(words1[:5])
#         if any(keyword in first_5_words_in_2 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words1[:4]).title().rstrip('.,')
#         else:
#             name_text = ' '.join(words1).title().rstrip('.,')
#     else:
#         first_5_words_in_1 = ' '.join(words[:5])
#         if any(keyword in first_5_words_in_1 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words[:4]).title().rstrip('.,')
#         else:
#             name_text = ' '.join(words).title().rstrip('.,')

#     return jsonify({
#         "name": name_text,
#         "mail": mail_text,
#         "phone": phone_text,
#         "skill1": skills_it,
#         "skill2": skills_non_it
#     })

###########################################################################################################


# def extract_text(file):
#     """
#     Extract text from PDF or DOCX files.
    
#     Parameters:
#         file (BytesIO): File-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     try:
#         file.seek(0)
#         header = file.read(4)
#         file.seek(0)
#         if header.startswith(b'%PDF'):
#             return extract_text_from_pdf(file)
#         elif header.startswith(b'PK\x03\x04'):
#             return extract_text_from_docx(file)
#         else:
#             return ""  # Unsupported file format
#     except Exception as e:
#         print(f"Error determining file type: {e}")
#         return ""

# def extract_text_from_pdf(file):
#     """
#     Extract text from a PDF file.
    
#     Parameters:
#         file (BytesIO): PDF file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         with fitz.open(stream=file, filetype="pdf") as doc:
#             for page in doc:
#                 text += page.get_text()
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#     return text

# def extract_text_from_docx(file):
#     """
#     Extract text from a DOCX file.
    
#     Parameters:
#         file (BytesIO): DOCX file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         doc = Document(file)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#     return text

# def extract_skills_from_resume(text, skills_list):
#     found_skills = [skill for skill in skills_list if skill.lower() in text.lower()]
#     return found_skills

# def extract_email(text):
#     email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#     email_matches = re.findall(email_regex, text)
#     return email_matches[0] if email_matches else "No email found"

# def extract_phone_number(text):
#     phone_regex = r'\+?\d[\d -]{8,12}\d'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[0] if phone_matches else "No phone number found"

# @app.route('/parse_resume', methods=['POST'])
# def parse_resume():
#     if 'resume' not in request.json:
#         return jsonify({"error": "No resume data provided"}), 400
    
#     data = request.json
#     resume_data = data['resume']
    
#     try:
#         decoded_resume = base64.b64decode(resume_data)
#     except Exception as e:
#         return jsonify({"error": "Invalid resume data"}), 400
    
#     resume_file = io.BytesIO(decoded_resume)
#     resume_text = extract_text(resume_file)
    
#     if not resume_text:
#         return jsonify({"error": "No text found in the resume data"}), 400

#     it_skills = [ 
#         'Data Analysis', 'Machine Learning', 'Communication', 'Project Management',
#         'Deep Learning', 'SQL', 'Tableau', 'C++', 'C', 'Front End Development', 'JAVA', 
#         'Java Full Stack', 'React JS', 'Node JS','Programming (Python, Java, C++)',
#         'Data Analysis and Visualization','Artificial Intelligence','Programming',
#         'Database Management (SQL)','Web Development (HTML, CSS, JavaScript)',
#         'Machine Learning and Artificial Intelligence','Network Administration',
#         'Software Development and Testing','Embedded Systems','CAD and 3D Modeling',
#         'HTML5', 'CSS3', 'Jquery', 'Bootstrap', 'XML', 'JSON', 'ABAP', 'SAPUI5',
#         'Agile Methodology', 'Frontend Development', 'Jira', 'Odata', 'BTP', 'Fiori Launchpad', 
#         'Python', 'JavaScript', 'HTML', 'CSS','React', 'Node.js', 'Django', 'Git', 'AWS',
#         'Linux','DevOps','Linear Regression','Logistic Regression','Decision Tree',
#         'SVM (Support Vector Machine)','Ensembles','Random Forest','Clustering',
#         'PCA (Principal Component Analysis)','K-means','Recommendation System',
#         'Market Basket Analysis','CNN','RNN','LSTM','Natural Language Processing',
#         'NLTK','LGBM','XGBoost','Transformers','Siamese network','BTYD (Buy Till You Die)',
#         'ML Ops Tools: Azure Synapse','Azure ML','Azure Databricks','ML flow','Airflow',
#         'Kubernetes','Dockers','Data Streaming – Kafka','Flask','LT Spice','Wireshark',
#         'Ansys Lumerical','Zemax OpticStudio','Xilinx Vivado','Google Collab','MATLAB'
#     ]
    
#     non_it_skills = [
#         'Communication Skills', 'Teamwork', 'Problem Solving', 'Time Management', 'Leadership',
#         'Creativity', 'Adaptability', 'Critical Thinking', 'Analytical Skills', 'Attention to Detail',
#         'Customer Service', 'Interpersonal Skills', 'Negotiation Skills', 'Project Management', 
#         'Presentation Skills', 'Research Skills', 'Organizational Skills', 'Multitasking',
#         'Decision Making', 'Emotional Intelligence', 'Conflict Resolution', 'Networking', 
#         'Strategic Planning', 'Public Speaking', 'Writing Skills', 'Sales Skills', 'Marketing', 
#         'Finance', 'Human Resources', 'Training and Development', 'Event Planning', 'Language Proficiency',
#         'Problem-Solving', 'Sales', 'Marketing', 'Financial Analysis', 'Customer Relationship Management (CRM)', 
#         'Quality Management', 'Supply Chain Management', 'Logistics', 'Health and Safety', 'Public Relations', 
#         'Social Media Management', 'Content Creation', 'Graphic Design', 'Video Editing', 'Photography', 
#         'Data Entry', 'Administrative Support', 'Customer Support', 'Teaching', 'Mentoring', 'Coaching', 
#         'Retail Management', 'Hospitality Management', 'Event Management', 'Creative Writing', 'Content Marketing', 
#         'Copywriting', 'Publications', 'Translation', 'Counseling', 'Fitness Instruction', 'Nutrition', 'Wellness', 
#         'Fashion Design', 'Interior Design', 'Artistic Skills', 'Music', 'Sports', 'Culinary Arts', 'Photography', 
#         'Videography', 'Project Coordination', 'Community Outreach', 'Volunteer Management', 'Fundraising', 
#         'Political Campaigning', 'Government Relations', 'Policy Analysis', 'Nonprofit Management', 'Grant Writing', 
#         'Fundraising', 'Event Planning', 'Real Estate', 'Property Management', 'Construction Management', 
#         'Facilities Management', 'Environmental Sustainability', 'Energy Management', 'Public Health', 
#         'Healthcare Administration', 'Nursing', 'Dental Hygiene', 'Pharmacy', 'Physical Therapy', 'Occupational Therapy', 
#         'Social Work', 'Child Care', 'Elderly Care', 'Counseling', 'Psychology', 'Sociology', 'Anthropology', 'Archaeology', 
#         'Geography', 'History', 'Political Science', 'Economics', 'Philosophy', 'Theology', 'Linguistics', 'Literature', 
#         'Creative Writing', 'Journalism', 'Broadcasting', 'Public Relations', 'Marketing', 'Advertising', 'Market Research', 
#         'Retail Sales', 'Wholesale Sales', 'Account Management', 'Client Relations', 'Customer Service', 'Conflict Resolution', 
#         'Presentation Skills', 'Public Speaking', 'Writing', 'Editing', 'Proofreading', 'Content Creation', 'Graphic Design', 
#         'Visual Merchandising', 'Retail Operations', 'Inventory Management', 'Supply Chain', 'Logistics', 'Quality Assurance', 
#         'Process Improvement', 'Project Management', 'Financial Planning', 'Budgeting', 'Financial Analysis', 'Bookkeeping', 
#         'Data Entry', 'Administrative Support', 'Executive Assistance', 'Time Management', 'Organizational Skills', 'Event Planning', 
#         'Event Coordination', 'Event Marketing', 'Catering', 'Venue Management', 'Wedding Planning', 'Trade Show Coordination', 
#         'Customer Service', 'Conflict Resolution', 'Problem-Solving', 'Decision Making', 'Team Collaboration', 'Leadership', 
#         'Supervision', 'Employee Training', 'Performance Management', 'Recruitment', 'Human Resources', 'Payroll Administration', 
#         'Employee Relations', 'Safety Compliance', 'Labor Relations', 'Legal Compliance', 'Contract Negotiation', 'Risk Management', 
#         'Policy Development', 'Quality Management', 'Process Improvement', 'Supply Chain Management', 'Logistics', 'Inventory Control', 
#         'Procurement', 'Distribution', 'Quality Assurance', 'Process Improvement', 'Product Development', 'Marketing Strategy', 
#         'Brand Management', 'Market Research', 'Public Relations', 'Social Media Management', 'Content Creation', 'Copywriting', 
#         'Email Marketing', 'Sales Strategy', 'Client Relationship Management', 'Sales Forecasting', 'Lead Generation', 
#         'Account Management', 'Customer Retention', 'Sales Presentations', 'Networking', 'Public Speaking', 'Team Collaboration', 
#         'Project Management', 'Client Communications', 'Technical Support', 'Troubleshooting', 'Network Administration', 
#         'Quality Assurance', 'Project Management', 'Technical Writing', 'Documentation', 'Research and Development', 'Innovation', 
#         'Problem-Solving', 'Critical Thinking', 'Attention to Detail', 'Collaboration', 'Time Management', 'Adaptability', 'Leadership', 
#         'Creativity', 'Analytical Skills', 'Data Analysis', 'Statistical Analysis', 'Mathematics', 'Physics', 'Chemistry', 'Biology', 
#         'Geology', 'Environmental Science', 'Meteorology', 'Agricultural Science', 'Animal Science', 'Food Science', 'Nutrition', 
#         'Dietetics', 'Physical Therapy', 'Occupational Therapy', 'Speech-Language Pathology', 'Nursing', 'Pharmacy', 'Dentistry', 
#         'Veterinary Medicine', 'Medical Research', 'Medical Writing', 'Clinical Trials', 'Epidemiology', 'Public Health', 
#         'Healthcare Administration', 'Health Informatics', 'Fitness Instruction', 'Nutrition Counseling', 'Wellness Coaching', 
#         'Yoga Instruction', 'Personal Training', 'Physical Education', 'Sports Coaching', 'Athletic Training', 'Recreation', 
#         'Dance Instruction', 'Music Instruction', 'Art Instruction', 'Photography', 'Video Editing', 'Graphic Design', 'Interior Design', 
#         'Fashion Design', 'Culinary Arts', 'Baking', 'Cooking', 'Restaurant Management', 'Hotel Management', 'Tourism', 'Event Planning', 
#         'Museum Management', 'Library Science', 'Archiving', 'Curatorial Work', 'Conservation', 'Environmental Science', 
#         'Sustainability', 'Renewable Energy', 'Climate Change', 'Environmental Policy', 'Wildlife Conservation', 'Forestry', 
#         'Natural Resource Management', 'Ecology', 'Geography', 'Urban Planning', 'Civil Engineering', 'Structural Engineering', 
#         'Transportation Engineering', 'Geotechnical Engineering', 'Environmental Engineering', 'Water Resources Engineering', 
#         'Surveying', 'Architecture', 'Landscape Architecture', 'Interior Design', 'Urban Design', 'Real Estate Development', 
#         'Property Management', 'Construction Management', 'Building Inspection', 'Facilities Management', 'Space Planning', 
#         'Urban Planning', 'Public Administration', 'Policy Analysis', 'Government Relations', 'Political Campaigning', 'Public Policy', 
#         'Economics', 'Finance', 'Accounting', 'Actuarial Science', 'MS Office','Powerpoint','ms word','ms excel' 
#     ]

#     extracted_skills = extract_skills_from_resume(resume_text, it_skills)
#     extracted_nonit_skills = extract_skills_from_resume(resume_text, non_it_skills)
#     non_it_skills = list(set(extracted_nonit_skills) - set(extracted_skills))

#     skills_it = ", ".join(extracted_skills) if extracted_skills else "No skills found"
#     skills_non_it = ", ".join(non_it_skills) if non_it_skills else "No skills found"

#     mail_text = extract_email(resume_text)
#     phone_text = extract_phone_number(resume_text)

#     first_line = resume_text.split('\n')[0]
#     words = first_line.split()
#     first_3_words = ' '.join(words[:3])

#     name_text = "No name found"
#     if "RESUME" in first_line or "Resume" in first_line or "BIODATA" in first_line or "BioData" in first_line or "biodata" in first_line:
#         second_line = resume_text.split('\n')[1] if len(resume_text.split('\n')) > 1 else ""
#         words1 = second_line.split()
#         first_5_words_in_2 = ' '.join(words1[:5])
#         if any(keyword in first_5_words_in_2 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words1[:4]).title()
#         else:
#             name_text = ' '.join(words1).title()
#     else:
#         first_5_words_in_1 = ' '.join(words[:5])
#         if any(keyword in first_5_words_in_1 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words[:4]).title()
#         else:
#             name_text = ' '.join(words).title()

#     return jsonify({
#         "name": name_text,
#         "mail": mail_text,
#         "phone": phone_text,
#         "skill1": skills_it,
#         "skill2": skills_non_it
#     })





# def extract_text(file_path):
#     """
#     Extract text from PDF or DOCX files.
    
#     Parameters:
#         file_path (str): Path to the file.
    
#     Returns:
#         str: Extracted text.
#     """
#     if file_path.endswith('.pdf'):
#         return extract_text_from_pdf(file_path)
#     elif file_path.endswith('.docx'):
#         return extract_text_from_docx(file_path)
#     else:
#         return ""  # Unsupported file format

# def extract_text_from_pdf(pdf_path):
#     """
#     Extract text from a PDF file.
    
#     Parameters:
#         pdf_path (str): Path to the PDF file.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         with fitz.open(pdf_path) as doc:
#             for page in doc:
#                 text += page.get_text()
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#     return text

# def extract_text_from_docx(docx_path):
#     """
#     Extract text from a DOCX file.
    
#     Parameters:
#         docx_path (str): Path to the DOCX file.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         doc = Document(docx_path)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#     return text

# def extract_skills_from_resume(text, skills_list):
#     # This function will search for skills in the text based on the given skills list
#     found_skills = [skill for skill in skills_list if skill.lower() in text.lower()]
#     return found_skills

# def extract_email(text):
#     # Regex to find email in text
#     email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#     email_matches = re.findall(email_regex, text)
#     return email_matches[0] if email_matches else "No email found"

# def extract_phone_number(text):
#     # Regex to find phone number in text
#     phone_regex = r'\+?\d[\d -]{8,12}\d'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[0] if phone_matches else "No phone number found"

# @app.route('/parse_resume', methods=['POST'])
# def parse_resume():
#     if 'resume' not in request.json:
#         return jsonify({"error": "No resume data provided"}), 400
    
#     data = request.json
#     resume_data = data['resume']
    
#     print("resume_data :",resume_data)

#     try:
#         decoded_resume = base64.b64decode(resume_data)
#     except Exception as e:
#         return jsonify({"error": "Invalid resume data"}), 400
    
#     resume_file = io.BytesIO(decoded_resume)
#     resume_text = extract_text(resume_file)
    
#     if not resume_text:
#         return jsonify({"error": "No text found in the resume data"}), 400

#     it_skills = [ 'Data Analysis', 'Machine Learning', 'Communication', 'Project Management',
#                     'Deep Learning', 'SQL', 'Tableau', 'C++', 'C', 'Front End Development', 'JAVA', 'Java Full Stack', 'React JS', 'Node JS','Programming (Python, Java, C++)',
#         'Data Analysis and Visualization','Artificial Intelligence','Programming'
#         'Database Management (SQL)',
#         'Web Development (HTML, CSS, JavaScript)',
#         'Machine Learning and Artificial Intelligence',
#         'Network Administration',
#         'Software Development and Testing',
#         'Embedded Systems',
#         'CAD and 3D Modeling',
#         'HTML5', 'CSS3'  ,'Jquery' ,'Bootstrap' ,'XML' ,'JSON' ,'ABAP'
#         ,'SAPUI5',
#         'Agile Methodology' ,'Frontend Development' ,'Jira' ,'Odata' ,'BTP' ,'Fiori Launchpad', 
#         'Python', 'JavaScript', 'HTML', 'CSS',
#         'React', 'Node.js', 'Django', 'Git', 'AWS', 'Linux','DevOps',
#         'Linear Regression',
#         'Logistic Regression',
#         'Decision Tree',
#         'SVM (Support Vector Machine)',
#         'Ensembles',
#         'Random Forest',
#         'Clustering',
#         'PCA (Principal Component Analysis)',
#         'K-means',
#         'Recommendation System',
#         'Market Basket Analysis',
#         'CNN',
#         'RNN',
#         'LSTM',
#         'Natural Language Processing',
#         'NLTK',
#         'LGBM',
#         'XGBoost',
#         'Transformers',
#         'Siamese network',
#         'BTYD (Buy Till You Die)',
#         'ML Ops Tools: Azure Synapse',
#         'Azure ML',
#         'Azure Databricks',
#         'ML flow',
#         'Airflow',
#         'Kubernetes',
#         'Dockers',
#         'Data Streaming – Kafka',
#         'Flask','LT Spice',
#         'Wireshark',
#         'Ansys Lumerical',
#         'Zemax OpticStudio',
#         'Xilinx Vivado',
#         'Google Collab',
#         'MATLAB'
#         ]
#     non_it_skills = ['Communication Skills', 'Teamwork', 'Problem Solving', 'Time Management', 'Leadership', 'Creativity', 'Adaptability', 'Critical Thinking', 'Analytical Skills', 'Attention to Detail', 'Customer Service', 'Interpersonal Skills', 'Negotiation Skills', 'Project Management', 'Presentation Skills', 'Research Skills', 'Organizational Skills', 'Multitasking', 'Decision Making', 'Emotional Intelligence', 'Conflict Resolution', 'Networking', 'Strategic Planning', 'Public Speaking', 'Writing Skills', 'Sales Skills', 'Marketing', 'Finance', 'Human Resources', 'Training and Development', 'Event Planning', 'Language Proficiency', 'Problem-Solving', 'Sales', 'Marketing', 'Financial Analysis', 'Customer Relationship Management (CRM)', 'Quality Management', 'Supply Chain Management', 'Logistics', 'Health and Safety', 'Public Relations', 'Social Media Management', 'Content Creation',
#                         'Graphic Design', 'Video Editing', 'Photography', 'Data Entry', 'Administrative Support', 'Customer Support', 'Teaching', 'Mentoring', 'Coaching', 'Retail Management', 'Hospitality Management', 'Event Management', 'Creative Writing', 'Content Marketing', 'Copywriting', 'Publications', 'Translation', 'Counseling', 'Fitness Instruction', 'Nutrition', 'Wellness', 'Fashion Design', 'Interior Design', 'Artistic Skills', 'Music', 'Sports', 'Culinary Arts', 'Photography', 'Videography', 'Project Coordination', 'Community Outreach', 'Volunteer Management', 'Fundraising', 'Political Campaigning', 'Government Relations', 'Policy Analysis', 'Nonprofit Management', 'Grant Writing', 'Fundraising', 'Event Planning', 'Real Estate', 'Property Management', 'Construction Management', 'Facilities Management', 'Environmental Sustainability', 'Energy Management', 'Public Health', 'Healthcare Administration', 
#                         'Nursing', 'Dental Hygiene', 'Pharmacy', 'Physical Therapy', 'Occupational Therapy', 'Social Work', 'Child Care', 'Elderly Care', 'Counseling', 'Psychology', 'Sociology', 'Anthropology', 'Archaeology', 'Geography', 'History', 'Political Science', 'Economics', 'Philosophy', 'Theology', 'Linguistics', 'Literature', 'Creative Writing', 'Journalism', 'Broadcasting', 'Public Relations', 'Marketing', 'Advertising', 'Market Research', 'Retail Sales', 'Wholesale Sales', 'Account Management', 'Client Relations', 'Customer Service', 'Conflict Resolution', 'Presentation Skills', 'Public Speaking', 'Writing', 'Editing', 'Proofreading', 'Content Creation', 'Graphic Design', 'Visual Merchandising', 'Retail Operations', 'Inventory Management', 'Supply Chain', 'Logistics', 'Quality Assurance', 'Process Improvement', 'Project Management', 'Financial Planning', 'Budgeting', 'Financial Analysis',
#                         'Bookkeeping', 'Data Entry', 'Administrative Support', 'Executive Assistance', 'Time Management', 'Organizational Skills', 'Event Planning', 'Event Coordination', 'Event Marketing', 'Catering', 'Venue Management', 'Wedding Planning', 'Trade Show Coordination', 'Customer Service', 'Conflict Resolution', 'Problem-Solving', 'Decision Making', 'Team Collaboration', 'Leadership', 'Supervision', 'Employee Training', 'Performance Management', 'Recruitment', 'Human Resources', 'Payroll Administration', 'Employee Relations', 'Safety Compliance', 'Labor Relations', 'Legal Compliance', 'Contract Negotiation', 'Risk Management', 'Policy Development', 'Quality Management', 'Process Improvement', 'Supply Chain Management', 'Logistics', 'Inventory Control', 'Procurement', 'Distribution', 'Quality Assurance', 'Process Improvement', 'Product Development', 'Marketing Strategy', 'Brand Management', 'Market Research', 
#                         'Public Relations', 'Social Media Management', 'Content Creation', 'Copywriting', 'Email Marketing', 'Sales Strategy', 'Client Relationship Management', 'Sales Forecasting', 'Lead Generation', 'Account Management', 'Customer Retention', 'Sales Presentations', 'Networking', 'Public Speaking', 'Team Collaboration', 'Project Management', 'Client Communications', 'Technical Support', 'Troubleshooting', 'Network Administration',   'Quality Assurance', 'Project Management', 'Technical Writing', 'Documentation',
#                         'Research and Development', 'Innovation', 'Problem-Solving', 'Critical Thinking', 'Attention to Detail', 'Collaboration', 'Time Management', 'Adaptability', 'Leadership', 'Creativity', 'Analytical Skills', 'Data Analysis', 'Statistical Analysis', 'Mathematics', 'Physics', 'Chemistry', 'Biology', 'Geology', 'Environmental Science', 'Meteorology', 'Agricultural Science', 'Animal Science', 'Food Science', 'Nutrition', 'Dietetics', 'Physical Therapy', 'Occupational Therapy', 'Speech-Language Pathology', 'Nursing', 'Pharmacy', 'Dentistry', 'Veterinary Medicine', 'Medical Research', 'Medical Writing', 'Clinical Trials', 'Epidemiology', 'Public Health', 'Healthcare Administration', 'Health Informatics', 'Fitness Instruction', 'Nutrition Counseling', 'Wellness Coaching', 'Yoga Instruction', 'Personal Training', 'Physical Education', 'Sports Coaching', 'Athletic Training', 'Recreation', 'Dance Instruction', 'Music Instruction',
#                         'Art Instruction', 'Photography', 'Video Editing', 'Graphic Design', 'Interior Design', 'Fashion Design', 'Culinary Arts', 'Baking', 'Cooking', 'Restaurant Management', 'Hotel Management', 'Tourism', 'Event Planning', 'Museum Management', 'Library Science', 'Archiving', 'Curatorial Work', 'Conservation', 'Environmental Science', 'Sustainability', 'Renewable Energy', 'Climate Change', 'Environmental Policy', 'Wildlife Conservation', 'Forestry', 'Natural Resource Management', 'Ecology', 'Geography', 'Urban Planning', 'Civil Engineering', 'Structural Engineering', 'Transportation Engineering', 'Geotechnical Engineering', 'Environmental Engineering', 'Water Resources Engineering', 'Surveying', 'Architecture', 'Landscape Architecture', 'Interior Design', 'Urban Design', 'Real Estate Development', 'Property Management', 'Construction Management', 'Building Inspection', 'Facilities Management', 'Space Planning', 'Urban Planning', 
#                         'Public Administration', 'Policy Analysis', 'Government Relations', 'Political Campaigning', 'Public Policy', 'Economics', 'Finance', 'Accounting', 'Actuarial Science' ,'MS Office','Powerpoint','ms word','ms excel' ]

#     extracted_skills = extract_skills_from_resume(resume_text, it_skills)
#     extracted_nonit_skills = extract_skills_from_resume(resume_text, non_it_skills)
#     non_it_skills = list(set(extracted_nonit_skills) - set(extracted_skills))

#     skills_it = ", ".join(extracted_skills) if extracted_skills else "No skills found"
#     skills_non_it = ", ".join(non_it_skills) if non_it_skills else "No skills found"

#     mail_text = extract_email(resume_text)
#     phone_text = extract_phone_number(resume_text)

#     first_line = resume_text.split('\n')[0]  # Extract the first line
#     words = first_line.split()
#     first_3_words = ' '.join(words[:3])

#     name_text = "No name found"
#     if "RESUME" in first_line or "Resume" in first_line or "BIODATA" in first_line or "BioData" in first_line or "biodata" in first_line:

#         print("if : entering")
        
#         second_line = resume_text.split('\n')[1] if len(resume_text.split('\n')) > 1 else ""
#         words1 = second_line.split()
#         first_5_words_in_2 = ' '.join(words1[:5])
#         if any(keyword in first_5_words_in_2 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words1[:4]).title()
#         else:
#             name_text = ' '.join(words1).title()
#     else:
#         first_5_words_in_1 = ' '.join(words[:5])
#         if any(keyword in first_5_words_in_1 for keyword in ["+91", "91", "@"]):
#             name_text = ' '.join(words[:4]).title()
#         else:
#             name_text = ' '.join(words).title()

#     return jsonify({
#         "name": name_text,
#         "mail": mail_text,
#         "phone": phone_text,
#         "skill1": skills_it,
#         "skill2": skills_non_it
#     })


if __name__ == '__main__':
    app.run(debug=True, host="0.0.0.0",port=5000)
